"""Whole-region front-matter classification pass (DD-54/DD-35).

The LLM call is mocked at the integration boundary (monkeypatch `complete`,
mirroring tests/integration/test_detect.py) — no live API call. Three layers are
exercised: `classify_frontmatter_region` (the AI pass in isolation),
`_apply_region` (the apply-and-guard contract), and `preview_docx` (the pipeline
wiring: deterministic-first, a single whole-region pass over the front matter, ON
by default, `ai=False` to skip)."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from backend.models.contract_tree import BlockClassification, ExtractedBlock, ParsedDocument
from backend.models.llm import CompletionResult, TokenUsage
from backend.services.import_ import classify as classify_mod
from backend.services.import_ import pipeline
from docx import Document


def _result(text: str) -> CompletionResult:
    return CompletionResult(text=text, usage=TokenUsage())


def _ambiguous_docx(dest: Path) -> None:
    """Front matter the deterministic rules cannot fully place: the title and the
    intro passage carry no keyword, so they default to neutral `recital` +
    `uncertain` (DD-54) — exactly what the whole-region pass resolves."""
    doc = Document()
    doc.add_paragraph("MASTER AGREEMENT")
    doc.add_paragraph(
        "This introductory passage defies every keyword rule the classifier knows about"
    )
    doc.add_paragraph("NOW, THEREFORE IT IS AGREED AS FOLLOWS:")
    doc.add_paragraph("1. Scope")
    doc.save(str(dest))


# --- classify_frontmatter_region: the AI pass in isolation -------------------


async def test_region_applies_per_block_roles(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> CompletionResult:
        return _result(
            '```json\n{"blocks": ['
            '{"order": 0, "role": "title"},'
            '{"order": 1, "role": "parties"},'
            '{"order": 2, "role": "recital"}]}\n```'
        )

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    out = await classify_mod.classify_frontmatter_region(
        {0: "TECHNOLOGY LICENSING AGREEMENT", 1: "BETWEEN A and B", 2: "WHEREAS ..."}
    )
    assert out == {0: "title", 1: "parties", 2: "recital"}


async def test_region_enforces_at_most_one_title(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> CompletionResult:
        return _result(
            '{"blocks": ['
            '{"order": 0, "role": "title"},'
            '{"order": 1, "role": "title"},'
            '{"order": 2, "role": "agreement_statement"}]}'
        )

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    out = await classify_mod.classify_frontmatter_region({0: "a", 1: "b", 2: "c"})
    assert out[0] == "title"
    assert 1 not in out  # second title dropped -> block keeps its deterministic role
    assert out[2] == "agreement_statement"


async def test_region_tolerates_parse_failure(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> CompletionResult:
        return _result("sorry, I can't classify this front matter")

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    out = await classify_mod.classify_frontmatter_region({0: "x"})
    assert out == {}  # graceful failure, no crash


async def test_region_rejects_off_taxonomy_role(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> CompletionResult:
        return _result('{"blocks": [{"order": 0, "role": "clause"}]}')  # not a region role

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    out = await classify_mod.classify_frontmatter_region({0: "looks operative"})
    assert out == {}  # validation failure -> treated as unresolved


async def test_region_skips_llm_when_empty(monkeypatch: Any) -> None:
    async def boom(**_kwargs: Any) -> str:
        raise AssertionError("no front matter -> no LLM call")

    monkeypatch.setattr(classify_mod, "complete", boom)

    assert await classify_mod.classify_frontmatter_region({}) == {}


# --- _apply_region: apply + the §12 export-exclusion guard --------------------


def _doc(*texts: str) -> ParsedDocument:
    blocks = [ExtractedBlock(order=i, kind="paragraph", text=t) for i, t in enumerate(texts)]
    return ParsedDocument(blocks=blocks, extracted_chars=1, ceiling_chars=1)


async def test_apply_region_clears_uncertain_and_applies_role(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> CompletionResult:
        return _result('{"blocks": [{"order": 0, "role": "title"}]}')

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    doc = _doc("TECHNOLOGY LICENSING AGREEMENT", "AGREED AS FOLLOWS:")
    classifications = {
        0: BlockClassification(role="recital", uncertain=True),
        1: BlockClassification(role="agreement_statement"),
    }

    await pipeline._apply_region(doc, classifications)

    assert classifications[0].role == "title"
    assert classifications[0].uncertain is False


async def test_apply_region_never_overrides_deterministic_drafting_note(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> CompletionResult:
        return _result('{"blocks": [{"order": 0, "role": "title"}]}')

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    doc = _doc("[CAM Notes: ...]", "AGREED AS FOLLOWS:")
    classifications = {
        0: BlockClassification(role="drafting_note"),
        1: BlockClassification(role="agreement_statement"),
    }

    await pipeline._apply_region(doc, classifications)

    assert classifications[0].role == "drafting_note"  # export guard: not retitled


async def test_apply_region_model_note_is_surfaced_not_silently_excluded(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> CompletionResult:
        return _result('{"blocks": [{"order": 0, "role": "drafting_note"}]}')

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    doc = _doc("[some bracketed aside the rules missed]", "AGREED AS FOLLOWS:")
    classifications = {
        0: BlockClassification(role="recital", uncertain=False),
        1: BlockClassification(role="agreement_statement"),
    }

    await pipeline._apply_region(doc, classifications)

    # Set to drafting_note BUT flagged uncertain for operator confirmation (DD-54):
    # never silently excluded from export on the model's say-so alone.
    assert classifications[0].role == "drafting_note"
    assert classifications[0].uncertain is True


# --- preview_docx: the pipeline wiring (ON by default) -----------------------


async def test_preview_ai_runs_region_pass(tmp_path: Path, monkeypatch: Any) -> None:
    fixture = tmp_path / "c.docx"
    _ambiguous_docx(fixture)
    calls: list[dict[str, Any]] = []

    async def fake_complete(**kwargs: Any) -> CompletionResult:
        calls.append(kwargs)
        return _result(
            '{"blocks": ['
            '{"order": 0, "role": "title"},'
            '{"order": 1, "role": "recital"},'
            '{"order": 2, "role": "agreement_statement"}]}'
        )

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    resp = await pipeline.preview_docx(fixture, ai=True)

    assert len(calls) == 1  # ONE whole-region call, not one per block
    assert calls[0]["caller"] == "import.classify_frontmatter_region"
    assert calls[0]["tier"] == "low"
    assert any(n.role == "title" for n in resp.nodes)
    assert sum(1 for n in resp.nodes if n.role == "title") == 1


async def test_preview_ai_false_skips_llm(tmp_path: Path, monkeypatch: Any) -> None:
    fixture = tmp_path / "c.docx"
    _ambiguous_docx(fixture)

    async def boom(**_kwargs: Any) -> str:
        raise AssertionError("LLM must not be called when ai=False")

    monkeypatch.setattr(classify_mod, "complete", boom)

    resp = await pipeline.preview_docx(fixture, ai=False)

    # Deterministic neutral default is kept; the block is still flagged for F04.
    assert any(n.role == "recital" and n.uncertain for n in resp.nodes)


async def test_preview_ai_parse_failure_leaves_deterministic_roles(
    tmp_path: Path, monkeypatch: Any
) -> None:
    fixture = tmp_path / "c.docx"
    _ambiguous_docx(fixture)

    async def fake_complete(**_kwargs: Any) -> CompletionResult:
        return _result("no idea, sorry")

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    resp = await pipeline.preview_docx(fixture, ai=True)

    # Unparseable -> blocks keep their deterministic recital role + uncertain.
    assert any(n.role == "recital" and n.uncertain for n in resp.nodes)
