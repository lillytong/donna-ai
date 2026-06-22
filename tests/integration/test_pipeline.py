"""Import orchestrator chains parse -> persist with the DB boundary mocked.

A synthetic .docx exercises the real read_docx -> build_tree -> tree_to_node_rows
chain; `insert_nodes`'s asyncpg connection is faked (no live database). The
optional entity-detection step is exercised with detect_entities monkeypatched
(the LLM boundary), mirroring tests/integration/test_detect.py.
"""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Any

from backend.models.extraction import Extraction
from backend.services.import_ import pipeline
from docx import Document


class _FakeConn:
    """Records inserted rows and hands back a generated id, like a RETURNING id."""

    def __init__(self) -> None:
        self.inserts: list[tuple[Any, ...]] = []

    async def fetchval(self, _sql: str, *args: Any) -> str:
        self.inserts.append(args)
        return str(uuid.uuid4())


def _build_docx(dest: Path) -> None:
    doc = Document()
    doc.add_paragraph("1. Definitions")
    doc.add_paragraph("In this agreement the following terms apply.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Royalty"
    table.cell(0, 1).text = "5%"
    doc.save(str(dest))


async def test_import_persists_parsed_nodes(tmp_path: Path) -> None:
    fixture = tmp_path / "c.docx"
    _build_docx(fixture)
    conn = _FakeConn()

    result = await pipeline.import_docx(conn, "contract-1", fixture)

    assert result.contract_id == "contract-1"
    assert result.node_count == len(conn.inserts) > 0  # every node was inserted
    assert result.root_count >= 1
    assert result.entity_candidates is None  # detection off by default; no LLM touched


async def test_optional_detection_runs_when_requested(tmp_path: Path, monkeypatch: Any) -> None:
    fixture = tmp_path / "c.docx"
    _build_docx(fixture)
    conn = _FakeConn()

    async def fake_detect(_clause_text: str) -> Extraction:
        return Extraction(defined_terms=[], cross_references=[], parameters=[])

    monkeypatch.setattr(pipeline, "detect_entities", fake_detect)

    result = await pipeline.import_docx(conn, "contract-1", fixture, detect=True)

    assert result.entity_candidates is not None
    # Only prose nodes with text are detected on; the table node is skipped.
    assert all(isinstance(v, Extraction) for v in result.entity_candidates.values())
    assert len(result.entity_candidates) >= 1
