"""Haiku residue pass for content-role classification (DD-54/DD-35).

The LLM call is mocked at the integration boundary (monkeypatch `complete`,
mirroring tests/integration/test_detect.py) — no live API call. Two layers are
exercised: `classify_residue` (the AI pass in isolation) and `preview_docx` (the
pipeline wiring: deterministic-first, AI over the uncertain front-matter, ON by
default, `ai=False` to skip).
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from backend.models.contract_tree import BlockClassification, ParsedDocument
from backend.services.import_ import classify as classify_mod
from backend.services.import_ import pipeline
from docx import Document


def _ambiguous_docx(dest: Path) -> None:
    """Front matter with one block the deterministic rules cannot place: it has no
    title/date/parties/recital keyword, so the classifier defaults it to a neutral
    `recital` + `uncertain` (DD-54) — exactly the residue the AI pass resolves."""
    doc = Document()
    doc.add_paragraph("MASTER AGREEMENT")
    doc.add_paragraph(
        "This introductory passage defies every keyword rule the classifier knows about"
    )
    doc.add_paragraph("NOW, THEREFORE IT IS AGREED AS FOLLOWS:")
    doc.add_paragraph("1. Scope")
    doc.save(str(dest))


# --- classify_residue: the AI pass in isolation -----------------------------


async def test_residue_takes_confident_role(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> str:
        return '```json\n{"role": "parties", "confident": true}\n```'

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    out = await classify_mod.classify_residue({3: "BETWEEN Acme Corp and Beta Ltd"})
    assert out == {3: "parties"}


async def test_residue_omits_low_confidence(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> str:
        return '{"role": "recital", "confident": false}'

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    out = await classify_mod.classify_residue({3: "an unclear block"})
    assert out == {}  # absent key -> caller keeps the deterministic uncertain flag


async def test_residue_tolerates_parse_failure(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> str:
        return "I'm sorry, I can't classify this block."

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    out = await classify_mod.classify_residue({3: "weird text"})
    assert out == {}  # graceful failure, no crash


async def test_residue_rejects_off_taxonomy_role(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> str:
        return '{"role": "clause", "confident": true}'  # not a front-matter role

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    out = await classify_mod.classify_residue({3: "looks operative"})
    assert out == {}  # validation failure -> treated as unresolved


# --- _apply_residue: the clear-uncertain contract ----------------------------


async def test_apply_residue_clears_classification_uncertain(monkeypatch: Any) -> None:
    async def fake_complete(**_kwargs: Any) -> str:
        return '{"role": "parties", "confident": true}'

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    doc = ParsedDocument(
        blocks=[classify_mod.ExtractedBlock(order=5, kind="paragraph", text="between-ish text")],
        extracted_chars=1,
        ceiling_chars=1,
    )
    classifications = {5: BlockClassification(role="recital", uncertain=True)}

    await pipeline._apply_residue(doc, classifications)

    assert classifications[5].role == "parties"
    assert classifications[5].uncertain is False


async def test_apply_residue_ignores_uncertain_operative_blocks(monkeypatch: Any) -> None:
    async def boom(**_kwargs: Any) -> str:
        raise AssertionError("operative residue must not reach the LLM")

    monkeypatch.setattr(classify_mod, "complete", boom)

    doc = ParsedDocument(
        blocks=[classify_mod.ExtractedBlock(order=0, kind="paragraph", text="clause text")],
        extracted_chars=1,
        ceiling_chars=1,
    )
    # An uncertain `clause` (no-boundary fallback) is NOT front-matter -> skipped.
    classifications = {0: BlockClassification(role="clause", uncertain=True)}

    await pipeline._apply_residue(doc, classifications)

    assert classifications[0].role == "clause"
    assert classifications[0].uncertain is True


# --- preview_docx: the pipeline wiring (ON by default) -----------------------


async def test_preview_ai_resolves_uncertain_frontmatter(tmp_path: Path, monkeypatch: Any) -> None:
    fixture = tmp_path / "c.docx"
    _ambiguous_docx(fixture)
    calls: list[dict[str, Any]] = []

    async def fake_complete(**kwargs: Any) -> str:
        calls.append(kwargs)
        return '{"role": "parties", "confident": true}'

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    resp = await pipeline.preview_docx(fixture, ai=True)

    assert calls  # the residue reached the LLM
    assert calls[0]["caller"] == "import.classify_role"
    assert calls[0]["tier"] == "low"
    # The ambiguous block carried the AI role; nothing is left labeled recital.
    assert any(n.role == "parties" for n in resp.nodes)
    assert not any(n.role == "recital" for n in resp.nodes)


async def test_preview_ai_false_skips_llm(tmp_path: Path, monkeypatch: Any) -> None:
    fixture = tmp_path / "c.docx"
    _ambiguous_docx(fixture)

    async def boom(**_kwargs: Any) -> str:
        raise AssertionError("LLM must not be called when ai=False")

    monkeypatch.setattr(classify_mod, "complete", boom)

    resp = await pipeline.preview_docx(fixture, ai=False)

    # Deterministic neutral default is kept; the block is still flagged for F04.
    assert any(n.role == "recital" and n.uncertain for n in resp.nodes)


async def test_preview_ai_parse_failure_leaves_deterministic_role(
    tmp_path: Path, monkeypatch: Any
) -> None:
    fixture = tmp_path / "c.docx"
    _ambiguous_docx(fixture)

    async def fake_complete(**_kwargs: Any) -> str:
        return "no idea, sorry"

    monkeypatch.setattr(classify_mod, "complete", fake_complete)

    resp = await pipeline.preview_docx(fixture, ai=True)

    # Unparseable -> the block keeps its deterministic recital role + uncertain.
    assert any(n.role == "recital" and n.uncertain for n in resp.nodes)
