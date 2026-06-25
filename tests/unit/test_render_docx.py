"""Renderer pieces (F14): numbering re-derivation, the DD-37 caps transform,
table rendering, and the clause/non-clause numbering split. No DB, no network."""

from __future__ import annotations

import io
from typing import Any

from backend.models.imports import StoredNode
from backend.services.export.render_docx import _plan, render_contract_docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def _node(
    node_id: str,
    parent_id: str | None,
    order_index: int,
    *,
    heading: str | None = None,
    body: str | None = None,
    table_data: list[list[str]] | None = None,
    content_type: str = "prose",
    role: str = "clause",
) -> StoredNode:
    return StoredNode(
        id=node_id,
        parent_id=parent_id,
        order_index=order_index,
        content_type=content_type,
        heading=heading,
        body=body,
        table_data=table_data,
        plain_text=heading or body or "",
        role=role,
    )


def _runs(paragraph: object) -> list[Any]:
    return list(paragraph.runs)  # type: ignore[attr-defined]


def _num_pr(paragraph: object) -> object | None:
    """The paragraph's w:numPr element (Word auto-numbering), or None if absent.
    numPr lives under pPr, so a direct child search on the paragraph misses it."""
    pPr = paragraph._p.find(qn("w:pPr"))  # type: ignore[attr-defined]
    return None if pPr is None else pPr.find(qn("w:numPr"))


def test_numbering_rederivation_from_tree_position() -> None:
    """DD-02: numbers are the decimal-outline path from tree position, never stored.
    Root clauses are 1,2…; children 1.1…; a depth-2 node carries two dots."""
    nodes = [
        _node("a", None, 100, heading="Definitions"),
        _node("b", "a", 100, body="Meaning of terms."),
        _node("c", None, 200, heading="Term"),
        _node("d", "c", 100, heading="Initial Term"),
        _node("e", "d", 100, body="Five years."),
    ]
    numbered = {n.id: num for n, num in _plan(nodes)}
    assert numbered == {"a": "1", "b": "1.1", "c": "2", "d": "2.1", "e": "2.1.1"}


def test_non_clause_nodes_are_unnumbered() -> None:
    """Front-matter (non-clause role) consumes no clause position and gets no number;
    the first real clause still numbers from 1 (DD-54)."""
    nodes = [
        _node("t", None, 100, heading="Master Agreement", role="title"),
        _node("p", None, 200, body="Between A and B.", role="parties"),
        _node("c", None, 300, heading="Definitions", role="clause"),
    ]
    numbered = {n.id: num for n, num in _plan(nodes)}
    assert numbered["t"] is None
    assert numbered["p"] is None
    assert numbered["c"] == "1"


def test_caps_renders_uppercase_but_preserves_stored_case() -> None:
    """DD-37 / §2.1: caps is the Word all-caps *display* property — the run text
    stays original-case (content boundary), Word renders it uppercase."""
    nodes = [_node("a", None, 100, heading="Confidentiality")]
    style_config = {"levels": {"0": {"bold": True, "caps": True}}}

    data = render_contract_docx(nodes, style_config)
    doc = Document(io.BytesIO(data))
    run = _runs(doc.paragraphs[0])[0]

    assert run.text == "Confidentiality"  # stored case untouched
    assert run.font.all_caps is True  # rendered uppercase via w:caps
    assert run.font.bold is True


def test_clause_paragraph_carries_numbering_non_clause_does_not() -> None:
    nodes = [
        _node("t", None, 100, heading="Master Agreement", role="title"),
        _node("c", None, 200, heading="Definitions", role="clause"),
    ]
    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))

    title_p, clause_p = doc.paragraphs[0], doc.paragraphs[1]
    assert (
        title_p._p.find(qn("w:pPr")) is None
        or title_p._p.find(  # no numbering
            qn("w:pPr")
        ).find(qn("w:numPr"))
        is None
    )
    num_pr = clause_p._p.find(qn("w:pPr")).find(qn("w:numPr"))
    assert num_pr is not None
    assert num_pr.find(qn("w:ilvl")).get(qn("w:val")) == "0"


def test_table_cells_render_faithfully() -> None:
    rows = [["Parameter", "Value"], ["Royalty", "5%"], ["Term", "5 years"]]
    nodes = [_node("tbl", None, 100, content_type="table", table_data=rows)]

    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))

    assert len(doc.tables) == 1
    table = doc.tables[0]
    rendered = [[cell.text for cell in row.cells] for row in table.rows]
    assert rendered == rows


def test_deeper_levels_inherit_nearest_font_size() -> None:
    """A depth-2 clause with no explicit level inherits the deepest defined level's
    font SIZE (the one style attribute still config-driven — bold/caps are now house
    rules), so rendering never fails on under-specified configs."""
    nodes = [
        _node("a", None, 100, heading="Article"),
        _node("b", "a", 100, heading="Section"),
        _node("c", "b", 100, body="Deep clause body."),
    ]
    style_config = {"levels": {"0": {"font_size_pt": 16}, "1": {"font_size_pt": 13}}}
    data = render_contract_docx(nodes, style_config)
    doc = Document(io.BytesIO(data))
    deep_run = _runs(doc.paragraphs[2])[0]
    assert deep_run.text == "Deep clause body."
    assert deep_run.font.size == Pt(13)  # size inherited from level 1


def test_text_with_native_enumerator_gets_no_auto_numbering() -> None:
    """A clause whose body already opens with its own enumerator (dotted decimal or
    "(a)") renders verbatim with NO w:numPr — auto-numbering would double the marker
    (DD-43). A clean clause with no leading marker still gets Word numbering."""
    nodes = [
        _node("a", None, 100, body="5.2.1 Pre-numbered clause."),
        _node("b", None, 200, body="(a) Parenthesised list item."),
        _node("c", None, 300, body="Clean clause with no marker."),
    ]
    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))
    dotted, paren, clean = doc.paragraphs[0], doc.paragraphs[1], doc.paragraphs[2]
    assert _num_pr(dotted) is None
    assert _num_pr(paren) is None
    assert _num_pr(clean) is not None


def test_heading_is_bold_but_caps_keys_on_source_not_bold() -> None:
    """A heading is bold (house style) but the uppercase transform is NOT inferred
    from bold — it keys on the source caps property only (DD-37, issue #2). With an
    empty style_config a bold mixed-case heading stays mixed-case (no w:caps); a
    genuinely-uppercase source heading still renders uppercase via its own text."""
    nodes = [
        _node("a", None, 100, heading="Fees"),
        _node("b", None, 200, heading="CONFIDENTIALITY"),
    ]
    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))
    mixed = _runs(doc.paragraphs[0])[0]
    upper = _runs(doc.paragraphs[1])[0]
    assert mixed.font.bold is True
    assert mixed.font.all_caps is not True
    assert mixed.text == "Fees"
    assert upper.text == "CONFIDENTIALITY"


def test_heading_caps_transform_applies_when_source_level_carries_caps() -> None:
    """The DD-37 uppercase transform still fires when the source caps property is
    present: a level with caps:true uppercases the heading at render while the stored
    text stays mixed-case (content integrity §2.1)."""
    nodes = [_node("a", None, 100, heading="Definitions")]
    data = render_contract_docx(nodes, {"levels": {"0": {"caps": True}}})
    doc = Document(io.BytesIO(data))
    run = _runs(doc.paragraphs[0])[0]
    assert run.font.bold is True
    assert run.font.all_caps is True
    assert run.text == "Definitions"


def test_appendix_title_is_centered_page_break_and_bold() -> None:
    """An appendix_title starts its own centred page (page break before) and renders
    bold (DD-37 house style)."""
    nodes = [_node("a", None, 100, heading="Appendix A", role="appendix_title")]
    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))
    paragraph = doc.paragraphs[0]
    run = _runs(paragraph)[0]
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph._p.find(qn("w:pPr")).find(qn("w:pageBreakBefore")) is not None
    assert run.font.bold is True


def test_title_is_centered_bold_and_12pt() -> None:
    """The contract title (role=title) renders centred, bold and 12pt (front matter),
    mirroring appendix-title alignment but WITHOUT a page break. Display-only: the
    stored text is unchanged."""
    nodes = [_node("t", None, 100, heading="Master Agreement", role="title")]
    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))
    paragraph = doc.paragraphs[0]
    run = _runs(paragraph)[0]
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph._p.find(qn("w:pPr")) is None or (
        paragraph._p.find(qn("w:pPr")).find(qn("w:pageBreakBefore")) is None
    )
    assert run.font.bold is True
    assert run.font.size == Pt(12)
    assert run.text == "Master Agreement"


def test_default_body_size_is_11pt_and_body_after_title_is_left_aligned() -> None:
    """House-style default body size is 11pt when no level overrides it, and a body
    paragraph after the title stays left-aligned (only the title centres)."""
    nodes = [
        _node("t", None, 100, heading="Master Agreement", role="title"),
        _node("b", None, 200, body="Some plain body text.", role="parties"),
    ]
    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))
    body_p = doc.paragraphs[1]
    body_run = _runs(body_p)[0]
    # default (None) alignment renders as left
    assert body_p.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)
    assert body_run.font.size == Pt(11)
    assert body_run.text == "Some plain body text."


def test_all_caps_span_in_body_renders_bold_inline() -> None:
    """An all-caps emphasis span (≥4 caps) in body text renders as its own bold run
    while the surrounding text stays regular (DD-37 inline house style)."""
    nodes = [_node("a", None, 100, body="The CONFIDENTIALITY clause applies.")]
    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))
    runs = _runs(doc.paragraphs[0])
    bolded = [r for r in runs if r.font.bold]
    assert [r.text for r in bolded] == ["CONFIDENTIALITY"]
    assert any(r.text == "The " and r.font.bold is False for r in runs)


def test_leading_defined_term_renders_bold_inline() -> None:
    """A leading quoted defined term ("Affiliate" means …) renders as a bold run."""
    nodes = [_node("a", None, 100, body="“Affiliate” means any controlled entity.")]
    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))
    bolded = [r for r in _runs(doc.paragraphs[0]) if r.font.bold]
    assert bolded[0].text == "“Affiliate”"


def test_indent_shares_first_two_levels_then_steps_in() -> None:
    """ilvl 0 and ilvl 1 share a left indent (0); ilvl 2 steps in once
    (max(0, ilvl-1) — DD-37: 14 / 14.1 flush, 14.1.1 indented)."""
    nodes = [
        _node("a", None, 100, body="Top clause."),
        _node("b", "a", 100, body="Sub clause."),
        _node("c", "b", 100, body="Sub-sub clause."),
    ]
    data = render_contract_docx(nodes, {})
    doc = Document(io.BytesIO(data))
    top, sub, deep = doc.paragraphs[0], doc.paragraphs[1], doc.paragraphs[2]
    assert top.paragraph_format.left_indent == Pt(0)
    assert sub.paragraph_format.left_indent == Pt(0)
    assert deep.paragraph_format.left_indent == Pt(18)
    assert deep.paragraph_format.left_indent > top.paragraph_format.left_indent
