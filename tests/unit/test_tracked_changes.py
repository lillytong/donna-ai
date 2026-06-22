"""count_tracked_changes surfaces <w:ins>/<w:del> for the clean-document guard (DD-46)."""

from __future__ import annotations

from pathlib import Path

from backend.services.import_.docx_reader import count_tracked_changes
from docx import Document
from docx.oxml import OxmlElement


def _run(text: str, tag: str = "w:t") -> OxmlElement:
    r = OxmlElement("w:r")
    t = OxmlElement(tag)
    t.text = text
    r.append(t)
    return r


def _build_tracked(dest: Path, *, insertions: int, deletions: int) -> None:
    doc = Document()
    p = doc.add_paragraph()
    for _ in range(insertions):
        ins = OxmlElement("w:ins")
        ins.append(_run("added text"))
        p._p.append(ins)
    for _ in range(deletions):
        dele = OxmlElement("w:del")
        dele.append(_run("removed text", tag="w:delText"))
        p._p.append(dele)
    doc.save(str(dest))


def test_counts_insertions_and_deletions(tmp_path: Path) -> None:
    fixture = tmp_path / "tracked.docx"
    _build_tracked(fixture, insertions=2, deletions=3)
    assert count_tracked_changes(fixture) == (2, 3)


def test_clean_document_reports_zero(tmp_path: Path) -> None:
    fixture = tmp_path / "clean.docx"
    doc = Document()
    doc.add_paragraph("A clean clause with no tracked changes.")
    doc.save(str(fixture))
    assert count_tracked_changes(fixture) == (0, 0)
