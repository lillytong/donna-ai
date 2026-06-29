"""Inline heading split: bold label + <w:br/> + body in one <w:p> -> two blocks.

Word Appendix paragraphs sometimes encode a bold label and body text in a single
<w:p>, separated by a soft line break (<w:br/>). Without this fix, _accept_all_text
concatenates them into a single string like "Licensing Fee10% of after-tax revenue...".
_split_at_line_break detects the <w:br/> and splits the paragraph into two blocks.

The <w:br/> is the discriminator: mixed bold/non-bold WITHOUT <w:br/> are inline
defined-term references and must NOT be split.
"""

from __future__ import annotations

from pathlib import Path

from backend.services.import_.docx_reader import read_docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


def _build_fixture(dest: Path) -> None:
    doc = Document()

    # 1. Normal paragraph with enough text to be body.
    doc.add_paragraph("1. This is a regular clause with enough text to be body.")

    # 2. Inline-heading paragraph: bold "Licensing Fee" + <w:br/> + plain body text.
    #    This is the bug case: one <w:p> encoding both a label and its body.
    p = doc.add_paragraph()
    p._p.clear()  # remove any default runs added by python-docx

    # Heading run: bold "Licensing Fee"
    heading_run = OxmlElement("w:r")
    rPr_h = OxmlElement("w:rPr")
    b_el = OxmlElement("w:b")
    rPr_h.append(b_el)
    heading_run.append(rPr_h)
    t_h = OxmlElement("w:t")
    t_h.text = "Licensing Fee"
    heading_run.append(t_h)
    p._p.append(heading_run)

    # Body run: soft line break + body text (non-bold).
    body_run = OxmlElement("w:r")
    rPr_b = OxmlElement("w:rPr")  # empty rPr — no <w:b/>
    body_run.append(rPr_b)
    # Add the soft line break (<w:br/> with no w:type attribute)
    br = etree.SubElement(body_run, qn("w:br"))  # noqa: F841 — side-effect: appends to body_run
    t_b = OxmlElement("w:t")
    t_b.text = "10% of after-tax revenue of the licensed facility per annum."
    body_run.append(t_b)
    p._p.append(body_run)

    # 3. Another normal paragraph.
    doc.add_paragraph("2. Another regular clause.")

    doc.save(str(dest))


def test_inline_heading_splits_into_two_blocks(tmp_path: Path) -> None:
    fixture = tmp_path / "inline_heading.docx"
    _build_fixture(fixture)

    parsed = read_docx(fixture)
    paras = [b for b in parsed.blocks if b.kind == "paragraph"]

    # The fixture has 3 source paragraphs but the inline-heading one splits into 2,
    # so we expect 4 paragraph blocks total.
    assert len(paras) == 4, (
        f"Expected 4 paragraph blocks (3 source + 1 split), got {len(paras)}. "
        f"Texts: {[b.text for b in paras]}"
    )

    texts = [b.text for b in paras]

    # The bold label must appear as its own block.
    assert "Licensing Fee" in texts, f"'Licensing Fee' heading block missing. Got: {texts}"

    # The body text must appear as its own block.
    body_expected = "10% of after-tax revenue of the licensed facility per annum."
    assert body_expected in texts, f"Body block missing. Got: {texts}"

    # The two must NOT be glued together.
    glued = [t for t in texts if "Licensing Fee" in t and "10%" in t]
    assert not glued, f"Heading and body are still glued: {glued}"

    # Normal clauses are still present and correct.
    assert any("1. This is a regular clause" in t for t in texts), (
        f"First normal clause missing. Got: {texts}"
    )
    assert any("2. Another regular clause" in t for t in texts), (
        f"Second normal clause missing. Got: {texts}"
    )
