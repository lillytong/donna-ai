"""Extraction is content-lossless and content-control-inclusive (DD-45).

Builds a synthetic .docx carrying the cases that break naive parsing:
  - a Word-auto-numbered clause (w:numPr)         -> numbering metadata
  - a tracked insertion + deletion                 -> accept-all-changes
  - a structured table                             -> rows, never flattened
  - a BLOCK-LEVEL content control (w:sdt)          -> text doc.paragraphs skips
"""

from __future__ import annotations

from pathlib import Path

from backend.services.import_.docx_reader import read_docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _run(text: str, tag: str = "w:t") -> OxmlElement:
    r = OxmlElement("w:r")
    t = OxmlElement(tag)
    t.text = text
    r.append(t)
    return r


def _build_fixture(dest: Path) -> None:
    doc = Document()

    # 1. Auto-numbered clause (w:numPr, level 1) with a literal "3.1" prefix too.
    p = doc.add_paragraph("3.1 Confidentiality terms apply.")
    pPr = OxmlElement("w:pPr")
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "1")
    numPr.append(ilvl)
    pPr.append(numPr)
    p._p.insert(0, pPr)

    # 2. Tracked insertion (kept) + deletion (dropped) in one paragraph.
    p2 = doc.add_paragraph()
    ins = OxmlElement("w:ins")
    ins.append(_run("USD 15/ton"))
    dele = OxmlElement("w:del")
    dele.append(_run("USD 10/ton", tag="w:delText"))
    p2._p.append(ins)
    p2._p.append(dele)

    # 3. Structured table.
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Parameter"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Royalty"
    table.cell(1, 1).text = "5%"

    # 4. Block-level content control wrapping a paragraph (the DD-45 case).
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    psdt = OxmlElement("w:p")
    psdt.append(_run("Counterparty Legal Name"))
    content.append(psdt)
    sdt.append(content)
    doc.element.body.append(sdt)

    doc.save(str(dest))


def test_extraction_is_lossless_and_content_control_inclusive(tmp_path: Path) -> None:
    fixture = tmp_path / "fixture.docx"
    _build_fixture(fixture)

    parsed = read_docx(fixture)
    paras = [b for b in parsed.blocks if b.kind == "paragraph"]
    texts = [b.text for b in paras]

    # DD-45: content-control text is captured, and flagged as such.
    cc = next(b for b in paras if "Counterparty Legal Name" in b.text)
    assert cc.in_content_control is True

    # Accept-all-changes: insertion kept, deletion dropped.
    assert any("USD 15/ton" in t for t in texts)
    assert all("USD 10/ton" not in t for t in texts)

    # Numbering metadata.
    numbered = next(b for b in paras if b.has_autonumber)
    assert numbered.list_level == 1
    assert numbered.literal_prefix == "3.1"

    # Table stays structured, never flattened.
    table = next(b for b in parsed.blocks if b.kind == "table")
    assert ["Royalty", "5%"] in (table.rows or [])

    # Coverage oracle: nothing silently lost.
    assert parsed.is_lossless, f"coverage only {parsed.coverage_pct:.1f}%"
