"""Round-trip + numbering oracle for block enumerated items (F03f, DD-98, §2.1).

A block enumerated item — "(a)"/"(A)"/"(i)" sitting as its own paragraph under a
clause — must, through the deterministic import spine and the renderer:
  - stay its own node (per-item addressable, not collapsed into one list block);
  - keep its literal marker as native body text (never stripped/re-derived);
  - NOT be decimal-renumbered (addressed as "1.1(b)", never "1.1.1");
  - round-trip byte-identical on content (import → render → re-extract).

The chain under test is the deterministic spine (no AI classify), so every node is
role=clause — exactly the case where, before F03f, the enumerated items were
decimal-renumbered, erasing the markers from addressing.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import pytest
from backend.models.imports import StoredNode
from backend.services.export.render_docx import render_contract_docx
from backend.services.import_.docx_reader import read_docx
from backend.services.import_.numbering import derive_numbers
from backend.services.import_.persist import tree_to_node_rows
from backend.services.import_.tree_builder import build_tree
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Three block enumerated schemes the spec names explicitly (DD-98).
_SCHEMES = {
    "lower_alpha": (
        "(a) both parties shall act;",
        "(b) no party shall object;",
        "(c) the term ends.",
    ),
    "upper_alpha": (
        "(A) the first recital applies;",
        "(B) the second applies;",
        "(C) the third applies.",
    ),
    "roman": ("(i) first condition holds;", "(ii) second holds;", "(iii) third holds."),
}


def _w(tag: str) -> str:
    return qn(f"w:{tag}")


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def _build_fixture(path: Path, items: tuple[str, ...]) -> None:
    """A clause lead-in (decimal auto-numbered) followed by block enumerated items
    as plain paragraphs whose literal marker lives in the text."""
    doc = Document()
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(_w("abstractNumId"), "20")
    for ilvl in range(2):
        lvl = OxmlElement("w:lvl")
        lvl.set(_w("ilvl"), str(ilvl))
        for tag, val in (("start", "1"), ("numFmt", "decimal")):
            el = OxmlElement(f"w:{tag}")
            el.set(_w("val"), val)
            lvl.append(el)
        txt = OxmlElement("w:lvlText")
        txt.set(_w("val"), ".".join(f"%{i + 1}" for i in range(ilvl + 1)) + ".")
        lvl.append(txt)
        abstract.append(lvl)
    numbering.insert(0, abstract)
    num = OxmlElement("w:num")
    num.set(_w("numId"), "21")
    ref = OxmlElement("w:abstractNumId")
    ref.set(_w("val"), "20")
    num.append(ref)
    numbering.append(num)

    def numbered(text: str, ilvl: int) -> None:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        il = OxmlElement("w:ilvl")
        il.set(_w("val"), str(ilvl))
        numPr.append(il)
        nid = OxmlElement("w:numId")
        nid.set(_w("val"), "21")
        numPr.append(nid)
        pPr.append(numPr)

    numbered("Definitions", 0)
    numbered("The following shall apply:", 1)
    for item in items:
        doc.add_paragraph(item)  # plain block paragraph — marker is in the text
    doc.save(str(path))


def _rows_to_stored(rows: list[Any]) -> list[StoredNode]:
    return [
        StoredNode(
            id=str(r.index),
            parent_id=str(r.parent_index) if r.parent_index is not None else None,
            order_index=r.order_index,
            content_type=r.content_type,
            heading=r.heading,
            body=r.body,
            table_data=r.table_data,
            plain_text=r.plain_text,
            role=r.role,
            has_placeholder=r.has_placeholder,
            enumerator_format=r.enumerator_format,
        )
        for r in rows
    ]


@pytest.mark.parametrize("scheme", list(_SCHEMES))
def test_block_enumerated_items_each_node_marker_kept_no_decimal(
    scheme: str, tmp_path: Path
) -> None:
    items = _SCHEMES[scheme]
    src = tmp_path / f"{scheme}.docx"
    _build_fixture(src, items)

    tree = build_tree(read_docx(src))
    rows = tree_to_node_rows(tree)
    numbers = derive_numbers(tree)

    # Locate the lead-in and its enumerated children.
    lead_in = next(n for n in tree.nodes if _norm(n.text) == "The following shall apply:")
    children = [n for n in tree.nodes if n.parent_index == lead_in.index]

    # 1. Each item is its own ordered child node (not collapsed).
    assert len(children) == len(items)
    # 2. Each child retains its literal marker as native text (verbatim).
    assert [_norm(c.text) for c in children] == [_norm(t) for t in items]
    # 3. No enumerated child is decimal-renumbered; the lead-in keeps its number.
    assert numbers.get(lead_in.index) == "1.1"
    for c in children:
        assert c.index not in numbers, f"enumerated item {c.text!r} got a decimal number"
    # The marker is body text, not a heading, and content_type stays prose (not list).
    enum_rows = [r for r in rows if r.parent_index == lead_in.index]
    assert all(r.content_type == "prose" for r in enum_rows)
    assert all(r.body is not None and r.heading is None for r in enum_rows)


@pytest.mark.parametrize("scheme", list(_SCHEMES))
def test_block_enumerated_roundtrip_byte_identical(scheme: str, tmp_path: Path) -> None:
    items = _SCHEMES[scheme]
    src = tmp_path / f"{scheme}.docx"
    _build_fixture(src, items)

    doc_a = read_docx(src)
    rows_a = tree_to_node_rows(build_tree(doc_a))

    data = render_contract_docx(_rows_to_stored(rows_a), {})
    rendered = tmp_path / "rendered.docx"
    rendered.write_bytes(data)

    doc_b = read_docx(rendered)
    tree_b = build_tree(doc_b)
    rows_b = tree_to_node_rows(tree_b)

    # Content integrity: same node count + per-node wording (whitespace normalised).
    assert len(rows_a) == len(rows_b), "node count drifted"
    for a, b in zip(rows_a, rows_b, strict=True):
        assert _norm(a.heading or a.body or "") == _norm(b.heading or b.body or "")
    # The marker text survives the render → re-extract round-trip exactly.
    para_texts = [_norm(b.text) for b in doc_b.blocks if b.kind == "paragraph"]
    for item in items:
        assert _norm(item) in para_texts, f"marker {item!r} lost in round-trip"
    # And the enumerated items stay unnumbered after the round-trip (symmetric skip).
    numbers_b = derive_numbers(tree_b)
    enum_b = [n for n in tree_b.nodes if _norm(n.text) in {_norm(t) for t in items}]
    assert enum_b and all(n.index not in numbers_b for n in enum_b)


# --- Auto-numbered (Word numPr) block enumerated items — the real-data shape (DD-99) ---

_AUTONUM = {
    "lowerLetter": ["(a)", "(b)", "(c)"],
    "upperLetter": ["(A)", "(B)", "(C)"],
    "lowerRoman": ["(i)", "(ii)", "(iii)"],
    "upperRoman": ["(I)", "(II)", "(III)"],
    # Parenthesised decimal "(1)(2)(3)" — detected on the parenthesised lvlText, not
    # on numFmt=decimal (the backbone clauses are decimal too). DD-99 amended / F03f.
    "decimal": ["(1)", "(2)", "(3)"],
}


def _build_autonum_fixture(path: Path, num_fmt: str) -> None:
    """A decimal lead-in clause + 3 children carrying Word auto-numbering (numPr) in
    `num_fmt` — the marker is generated by Word, absent from the run text (the real
    JVA shape). Mirrors how the real gitignored JVA fixture encodes its `(a)(b)(c)` items."""
    doc = Document()
    numbering = doc.part.numbering_part.element

    # Decimal backbone abstractNum (id 30) + alpha/roman abstractNum (id 31).
    dec = OxmlElement("w:abstractNum")
    dec.set(_w("abstractNumId"), "30")
    for ilvl in range(2):
        lvl = OxmlElement("w:lvl")
        lvl.set(_w("ilvl"), str(ilvl))
        for tag, val in (("start", "1"), ("numFmt", "decimal")):
            el = OxmlElement(f"w:{tag}")
            el.set(_w("val"), val)
            lvl.append(el)
        t = OxmlElement("w:lvlText")
        t.set(_w("val"), ".".join(f"%{i + 1}" for i in range(ilvl + 1)) + ".")
        lvl.append(t)
        dec.append(lvl)
    numbering.insert(0, dec)
    enum_abs = OxmlElement("w:abstractNum")
    enum_abs.set(_w("abstractNumId"), "31")
    lvl = OxmlElement("w:lvl")
    lvl.set(_w("ilvl"), "0")
    for tag, val in (("start", "1"), ("numFmt", num_fmt)):
        el = OxmlElement(f"w:{tag}")
        el.set(_w("val"), val)
        lvl.append(el)
    t = OxmlElement("w:lvlText")
    t.set(_w("val"), "(%1)")
    lvl.append(t)
    enum_abs.append(lvl)
    numbering.insert(0, enum_abs)
    for num_id, abs_id in (("30", "30"), ("31", "31")):
        num = OxmlElement("w:num")
        num.set(_w("numId"), num_id)
        ref = OxmlElement("w:abstractNumId")
        ref.set(_w("val"), abs_id)
        num.append(ref)
        numbering.append(num)

    def numbered(text: str, num_id: str, ilvl: int) -> None:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        il = OxmlElement("w:ilvl")
        il.set(_w("val"), str(ilvl))
        numPr.append(il)
        nid = OxmlElement("w:numId")
        nid.set(_w("val"), num_id)
        numPr.append(nid)
        pPr.append(numPr)

    numbered("Definitions", "30", 0)
    numbered("The following shall apply", "30", 1)  # decimal lead-in clause
    for body in (
        "first obligation applies",
        "second obligation applies",
        "third obligation applies",
    ):
        numbered(body, "31", 0)  # auto-numbered enumerated item — NO marker in text
    doc.save(str(path))


@pytest.mark.parametrize("fmt", list(_AUTONUM))
def test_autonumbered_items_captured_and_marker_derived(fmt: str, tmp_path: Path) -> None:
    from backend.services.import_.numbering import derive_enumerators, derive_numbers

    src = tmp_path / f"auto_{fmt}.docx"
    _build_autonum_fixture(src, fmt)

    tree = build_tree(read_docx(src))
    enum_nodes = [n for n in tree.nodes if n.enumerated and n.enumerator_format == fmt]
    # 3 items detected as auto-numbered enumerated, format captured, NO marker in body.
    assert len(enum_nodes) == 3
    assert all("(" not in (n.text or "") for n in enum_nodes), "marker leaked into body text"
    # Marker derived from position; NOT decimal-numbered.
    markers = derive_enumerators(tree)
    nums = derive_numbers(tree)
    assert sorted(markers[n.index] for n in enum_nodes) == sorted(_AUTONUM[fmt])
    assert all(n.index not in nums for n in enum_nodes)


@pytest.mark.parametrize("fmt", list(_AUTONUM))
def test_autonumbered_roundtrip_stable(fmt: str, tmp_path: Path) -> None:
    from backend.services.import_.numbering import derive_enumerators, derive_numbers

    src = tmp_path / f"auto_{fmt}.docx"
    _build_autonum_fixture(src, fmt)

    rows_a = tree_to_node_rows(build_tree(read_docx(src)))
    data = render_contract_docx(_rows_to_stored(rows_a), {})
    rendered = tmp_path / "rendered.docx"
    rendered.write_bytes(data)

    tree_b = build_tree(read_docx(rendered))
    rows_b = tree_to_node_rows(tree_b)

    # Content stable: same node count + per-node wording (markers never in the text).
    assert len(rows_a) == len(rows_b)
    for a, b in zip(rows_a, rows_b, strict=True):
        assert _norm(a.heading or a.body or "") == _norm(b.heading or b.body or "")
    # Enumerated items survive as auto-numbered, format recaptured, marker re-derived,
    # still NOT decimal — the round-trip is symmetric (numPr → numPr, never stored text).
    enum_b = [n for n in tree_b.nodes if n.enumerator_format == fmt]
    assert len(enum_b) == 3
    markers_b = derive_enumerators(tree_b)
    nums_b = derive_numbers(tree_b)
    assert sorted(markers_b[n.index] for n in enum_b) == sorted(_AUTONUM[fmt])
    assert all(n.index not in nums_b for n in enum_b)


# --- Render-XML defects: per-list restart + depth indent (BUG A / BUG B) ---


def _numbering_root(data: bytes, tmp_path: Path):
    rendered = tmp_path / "rendered.docx"
    rendered.write_bytes(data)
    return Document(str(rendered)).part.numbering_part.element


def _num_to_abstract(numbering: Any) -> dict[str, str]:
    """numId -> abstractNumId from the rendered numbering part."""
    out: dict[str, str] = {}
    for num in numbering.findall(_w("num")):
        ref = num.find(_w("abstractNumId"))
        if ref is not None:
            out[num.get(_w("numId"))] = ref.get(_w("val"))
    return out


def _enum_para_numpr(data: bytes, tmp_path: Path) -> list[tuple[str, str, int]]:
    """(text, numId, left-indent-twips) for each rendered paragraph that carries a
    numPr — only the enum paragraphs do here (front-matter clauses use a different
    numId; we key on the body text the test plants)."""
    rendered = tmp_path / "rendered.docx"
    rendered.write_bytes(data)
    body = Document(str(rendered)).part.document.element.body
    out: list[tuple[str, str, int]] = []
    for p in body.findall(_w("p")):
        pPr = p.find(_w("pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(_w("numPr"))
        if numPr is None:
            continue
        nid = numPr.find(_w("numId"))
        ind = pPr.find(_w("ind"))
        text = "".join(t.text or "" for t in p.iter(_w("t")))
        left = int(ind.get(_w("left"))) if ind is not None and ind.get(_w("left")) else 0
        out.append((_norm(text), nid.get(_w("val")) if nid is not None else "", left))
    return out


def test_separate_enum_lists_restart_distinct_abstractnums(tmp_path: Path) -> None:
    """BUG A: two separate lowerLetter lists under different parents must NOT share one
    counter. Each list-run gets its own numId AND its own abstractNum, so Word restarts
    each at `(a)` instead of continuing ((a)(b) then (c)(d)…)."""
    nodes = [
        StoredNode(id="A", parent_id=None, order_index=0, content_type="prose", body="Clause A"),
        StoredNode(
            id="a1",
            parent_id="A",
            order_index=0,
            content_type="prose",
            body="first under A",
            enumerator_format="lowerLetter",
        ),
        StoredNode(
            id="a2",
            parent_id="A",
            order_index=1,
            content_type="prose",
            body="second under A",
            enumerator_format="lowerLetter",
        ),
        StoredNode(id="B", parent_id=None, order_index=1, content_type="prose", body="Clause B"),
        StoredNode(
            id="b1",
            parent_id="B",
            order_index=0,
            content_type="prose",
            body="first under B",
            enumerator_format="lowerLetter",
        ),
        StoredNode(
            id="b2",
            parent_id="B",
            order_index=1,
            content_type="prose",
            body="second under B",
            enumerator_format="lowerLetter",
        ),
        StoredNode(
            id="b3",
            parent_id="B",
            order_index=2,
            content_type="prose",
            body="third under B",
            enumerator_format="lowerLetter",
        ),
    ]
    data = render_contract_docx(nodes, {})
    numbering = _numbering_root(data, tmp_path)
    num_to_abs = _num_to_abstract(numbering)

    enum = _enum_para_numpr(data, tmp_path)
    list_a = [e for e in enum if "under A" in e[0]]
    list_b = [e for e in enum if "under B" in e[0]]
    assert len(list_a) == 2 and len(list_b) == 3

    numid_a = {e[1] for e in list_a}
    numid_b = {e[1] for e in list_b}
    assert len(numid_a) == 1 and len(numid_b) == 1, "items within a list share one numId"
    assert numid_a != numid_b, "the two lists must use distinct numIds (independent runs)"

    abs_a = num_to_abs[next(iter(numid_a))]
    abs_b = num_to_abs[next(iter(numid_b))]
    assert abs_a != abs_b, "each list-run must own a distinct abstractNum so its counter restarts"


def test_enum_indent_scales_with_tree_depth(tmp_path: Path) -> None:
    """BUG B: a nested enum item steps in. A `(b)` at depth 1 and an `(A)` nested under
    it at depth 2 must carry a paragraph w:ind whose left grows with depth (the numbered
    `(A)` indents past the `(b)` it nests under)."""
    nodes = [
        StoredNode(id="C", parent_id=None, order_index=0, content_type="prose", body="Clause C"),
        StoredNode(
            id="b",
            parent_id="C",
            order_index=0,
            content_type="prose",
            body="outer b item",
            enumerator_format="lowerLetter",
        ),
        StoredNode(
            id="A",
            parent_id="b",
            order_index=0,
            content_type="prose",
            body="nested A item",
            enumerator_format="upperLetter",
        ),
    ]
    data = render_contract_docx(nodes, {})
    enum = {e[0]: e for e in _enum_para_numpr(data, tmp_path)}
    outer = enum[_norm("outer b item")]
    nested = enum[_norm("nested A item")]
    assert nested[2] > outer[2], "nested enum item must have a larger left indent than its parent"
    # ilvl pinned at 0 (round-trip constraint): only the paragraph w:ind carries depth.
    rendered = tmp_path / "rendered.docx"
    rendered.write_bytes(data)
    body = Document(str(rendered)).part.document.element.body
    for p in body.findall(_w("p")):
        pPr = p.find(_w("pPr"))
        numPr = pPr.find(_w("numPr")) if pPr is not None else None
        if numPr is None:
            continue
        ilvl = numPr.find(_w("ilvl"))
        assert ilvl is None or ilvl.get(_w("val")) == "0"


def test_nested_sublist_does_not_split_parent_run(tmp_path: Path) -> None:
    """BUG C (clause 14.2): a lowerLetter list `(a)(b)(c)` whose FIRST item `(a)` carries
    its OWN upperLetter sub-list `(A)(B)` must stay ONE run. The sub-list interleaves in
    DFS between `(a)` and `(b)`, but the three siblings share a single numId so Word
    numbers them a/b/c (not a, then a fresh a/b). The sub-list owns a different numId."""
    from backend.services.export.render_docx import _plan, _plan_enum_numbering

    nodes = [
        StoredNode(id="D", parent_id=None, order_index=0, content_type="prose", body="Clause D"),
        StoredNode(
            id="a",
            parent_id="D",
            order_index=0,
            content_type="prose",
            body="first sibling",
            enumerator_format="lowerLetter",
        ),
        StoredNode(
            id="aA",
            parent_id="a",
            order_index=0,
            content_type="prose",
            body="sub one",
            enumerator_format="upperLetter",
        ),
        StoredNode(
            id="aB",
            parent_id="a",
            order_index=1,
            content_type="prose",
            body="sub two",
            enumerator_format="upperLetter",
        ),
        StoredNode(
            id="b",
            parent_id="D",
            order_index=1,
            content_type="prose",
            body="second sibling",
            enumerator_format="lowerLetter",
        ),
        StoredNode(
            id="c",
            parent_id="D",
            order_index=2,
            content_type="prose",
            body="third sibling",
            enumerator_format="lowerLetter",
        ),
    ]
    node_numid, runs = _plan_enum_numbering(_plan(nodes))

    # The three lowerLetter siblings share ONE numId (one run, not split by the sub-list).
    sibling_numids = {node_numid["a"], node_numid["b"], node_numid["c"]}
    assert len(sibling_numids) == 1, "the outer (a)(b)(c) list was split into two runs"
    # The (A)(B) sub-list uses a DIFFERENT numId.
    sub_numids = {node_numid["aA"], node_numid["aB"]}
    assert len(sub_numids) == 1
    assert sibling_numids != sub_numids, "sub-list must own a distinct numId from its parent list"
    # Exactly one lowerLetter run for that parent (plus one upperLetter run for the sub).
    lower_runs = [r for r in runs if r[2] == "lowerLetter"]
    upper_runs = [r for r in runs if r[2] == "upperLetter"]
    assert len(lower_runs) == 1 and len(upper_runs) == 1
