"""Faithful .docx content extraction for the import spine (Phase 0).

Promotes the validated spike-#1 extraction (accept-all-changes text, w:numPr
numbering) and closes its one defect: block-level content controls (`w:sdt`)
whose paragraphs `python-docx`'s `doc.paragraphs` silently skips, losing
fill-in field text (party names, dates, placeholders) — DD-45.

Extraction is accept-all-changes: insertions (`w:ins`) kept, deletions
(`w:del`/`w:delText`) dropped. Formatting is intentionally not read here; it is
rebuilt from style_config on export (DD-43). Content only.
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from backend.models.contract_tree import ExtractedBlock, ParsedDocument

W_P = qn("w:p")
W_TBL = qn("w:tbl")
W_SDT = qn("w:sdt")
W_SDTCONTENT = qn("w:sdtContent")
W_T = qn("w:t")
W_INS = qn("w:ins")
W_DEL = qn("w:del")
W_PPR = qn("w:pPr")
W_NUMPR = qn("w:numPr")
W_ILVL = qn("w:ilvl")
W_NUMID = qn("w:numId")
W_VAL = qn("w:val")
W_TR = qn("w:tr")
W_TC = qn("w:tc")
W_PSTYLE = qn("w:pStyle")
W_OUTLINELVL = qn("w:outlineLvl")
W_STYLE = qn("w:style")
W_TYPE = qn("w:type")
W_STYLEID = qn("w:styleId")
W_BASEDON = qn("w:basedOn")
W_NUM = qn("w:num")
W_ABSTRACTNUM = qn("w:abstractNum")
W_ABSTRACTNUMID = qn("w:abstractNumId")
W_STYLELINK = qn("w:styleLink")
W_NUMSTYLELINK = qn("w:numStyleLink")
W_R = qn("w:r")
W_RPR = qn("w:rPr")
W_B = qn("w:b")
W_BR = qn("w:br")
W_LVL = qn("w:lvl")
W_NUMFMT = qn("w:numFmt")
W_DRAWING = qn("w:drawing")
WP_INLINE = qn("wp:inline")
WP_ANCHOR = qn("wp:anchor")
WP_EXTENT = qn("wp:extent")
A_BLIP = qn("a:blip")
# r:embed is an attribute in the Office relationships namespace.
R_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
_MIME_BY_EXT: dict[str, str] = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".bmp": "image/bmp",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
    ".emf": "image/emf",
    ".wmf": "image/wmf",
}

# Word's outlineLvl=9 is the "body text" sentinel, not a heading level (0-8).
_BODY_OUTLINE_LEVEL = 9


class _Numbering:
    """Resolved numbering context for one document: style-inherited (ilvl, numId,
    outlineLvl) per paragraph styleId (basedOn chains followed), and numId ->
    canonical abstractNumId (numStyleLink/styleLink indirection collapsed). Built
    once per import from styles.xml + numbering.xml; empty maps when those parts
    are absent (synthetic docs), in which case only direct numPr is read."""

    def __init__(self, style_index: dict[str, _StyleNum], abstract_map: dict[int, int], bullet_abs: set[int] | None = None) -> None:
        self._style_index = style_index
        self._abstract_map = abstract_map
        self._bullet_abs: set[int] = bullet_abs if bullet_abs is not None else set()

    def style(self, style_id: str | None) -> _StyleNum:
        if style_id is None:
            return (None, None, None)
        return self._style_index.get(style_id, (None, None, None))

    def abstract_of(self, num_id: int | None) -> int | None:
        if num_id is None:
            return None
        return self._abstract_map.get(num_id)

    def is_bullet(self, abstract_num_id: int | None) -> bool:
        if abstract_num_id is None:
            return False
        return abstract_num_id in self._bullet_abs


# (ilvl, numId, outlineLvl) carried (directly or by inheritance) by a paragraph style.
_StyleNum = tuple[int | None, int | None, int | None]


def _int_val(parent: Any, tag: str) -> int | None:
    el = parent.find(tag)
    if el is None:
        return None
    raw = el.get(W_VAL)
    return int(raw) if raw is not None else None


def _numpr_levels(numpr: Any) -> tuple[int | None, int | None]:
    """(ilvl, numId) read off a w:numPr element."""
    return _int_val(numpr, W_ILVL), _int_val(numpr, W_NUMID)


def _build_style_index(styles_root: Any) -> dict[str, _StyleNum]:
    """styleId -> resolved (ilvl, numId, outlineLvl), nearest definition winning
    over a basedOn ancestor. Only paragraph styles participate."""
    # styleId -> (based_on, own ilvl, own numId, own outlineLvl) before inheritance.
    raw: dict[str, tuple[str | None, int | None, int | None, int | None]] = {}
    for st in styles_root.findall(W_STYLE):
        if st.get(W_TYPE) != "paragraph":
            continue
        sid = st.get(W_STYLEID)
        if sid is None:
            continue
        pPr = st.find(W_PPR)
        ilvl = numid = outline = None
        if pPr is not None:
            numpr = pPr.find(W_NUMPR)
            if numpr is not None:
                ilvl, numid = _numpr_levels(numpr)
            outline = _int_val(pPr, W_OUTLINELVL)
        based_on = st.find(W_BASEDON)
        raw[sid] = (based_on.get(W_VAL) if based_on is not None else None, ilvl, numid, outline)

    resolved: dict[str, _StyleNum] = {}

    def resolve(sid: str | None, stack: frozenset[str]) -> _StyleNum:
        if sid is None or sid not in raw or sid in stack:
            return (None, None, None)
        if sid in resolved:
            return resolved[sid]
        based_on, ilvl, numid, outline = raw[sid]
        p_ilvl, p_numid, p_outline = resolve(based_on, stack | {sid})
        out: _StyleNum = (
            ilvl if ilvl is not None else p_ilvl,
            numid if numid is not None else p_numid,
            outline if outline is not None else p_outline,
        )
        resolved[sid] = out
        return out

    for sid in raw:
        resolve(sid, frozenset())
    return resolved


def _build_abstract_map(numbering_root: Any) -> dict[int, int]:
    """numId -> canonical abstractNumId. A numId references an abstractNumId; an
    abstractNum that is a numStyleLink shell (it delegates to a paragraph style)
    is collapsed onto the abstractNum that *defines* that style (styleLink), so
    two numIds backing one outline group together (DD-36)."""
    num_to_abs: dict[int, int] = {}
    for num in numbering_root.findall(W_NUM):
        nid = num.get(W_NUMID)
        abs_el = num.find(W_ABSTRACTNUMID)
        if nid is not None and abs_el is not None and abs_el.get(W_VAL) is not None:
            num_to_abs[int(nid)] = int(abs_el.get(W_VAL))

    style_to_abs: dict[str, int] = {}
    num_style_link: dict[int, str] = {}
    for an in numbering_root.findall(W_ABSTRACTNUM):
        aid_raw = an.get(W_ABSTRACTNUMID)
        if aid_raw is None:
            continue
        aid = int(aid_raw)
        sl = an.find(W_STYLELINK)
        if sl is not None and sl.get(W_VAL) is not None:
            style_to_abs[sl.get(W_VAL)] = aid
        nsl = an.find(W_NUMSTYLELINK)
        if nsl is not None and nsl.get(W_VAL) is not None:
            num_style_link[aid] = nsl.get(W_VAL)

    def canon(aid: int) -> int:
        seen: set[int] = set()
        while aid in num_style_link and aid not in seen:
            seen.add(aid)
            target = style_to_abs.get(num_style_link[aid])
            if target is None or target == aid:
                break
            aid = target
        return aid

    return {nid: canon(aid) for nid, aid in num_to_abs.items()}


def _build_bullet_set(numbering_root: Any) -> set[int]:
    """abstractNumIds whose primary level format is 'bullet'."""
    bullet_abs: set[int] = set()
    for an in numbering_root.findall(W_ABSTRACTNUM):
        aid_raw = an.get(W_ABSTRACTNUMID)
        if aid_raw is None:
            continue
        aid = int(aid_raw)
        for lvl in an.findall(W_LVL):
            nf = lvl.find(W_NUMFMT)
            if nf is not None and nf.get(W_VAL) == "bullet":
                bullet_abs.add(aid)
                break
    return bullet_abs


# Typed clause-number prefixes the parser must recognise: "3.", "3.1", "(a)", "(iv)".
_LITERAL_NUM = re.compile(r"^\s*(\d+(?:\.\d+)*\.?|\([a-z]+\)|\([ivx]+\))\s+\S")


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def _run_is_bold(r: Any) -> bool:
    """True when the run carries <w:b/> not explicitly cleared (w:val=0/false)."""
    rPr = r.find(W_RPR)
    if rPr is None:
        return False
    b = rPr.find(W_B)
    if b is None:
        return False
    val = b.get(W_VAL)
    return val not in ("0", "false")


def _split_at_line_break(p: Any) -> tuple[str, str] | None:
    """Detect paragraphs where a <w:br/> separates a bold label from plain body text.

    Word's Appendix/Annexure inline headings use a soft line break to visually separate
    a bold label ("Licensing Fee") from body text within the same <w:p>.
    _accept_all_text ignores <w:br/> and concatenates them -> "Licensing Fee10% of...".

    The <w:br/> is the discriminator: paragraphs with mixed bold/non-bold runs but NO
    <w:br/> are inline defined-term references and must NOT be split.

    Returns (heading_text, body_text) or None if pattern not present.
    Guards: heading <= 80 chars, no sentence-terminal punctuation; body >= 10 chars;
    pre-br content must be (at least partly) bold; post-br must not be all-bold."""
    run_entries: list[tuple[str, bool, bool]] = []  # (text, is_bold, has_br)
    for child in p:
        tag = child.tag
        if tag == W_DEL or tag == W_PPR:
            continue
        if tag == W_R:
            t = _norm(_accept_all_text(child))
            run_entries.append((t, _run_is_bold(child), child.find(W_BR) is not None))
        elif tag == W_INS:
            for r in child.findall(W_R):
                t = _norm(_accept_all_text(r))
                run_entries.append((t, _run_is_bold(r), r.find(W_BR) is not None))

    br_idx = next((i for i, (_, _, has_br) in enumerate(run_entries) if has_br), None)
    if br_idx is None or br_idx == 0:
        return None

    before = run_entries[:br_idx]
    after = run_entries[br_idx:]

    heading_text = "".join(t for t, _, _ in before).strip()
    body_text = "".join(t for t, _, _ in after).strip()

    if not heading_text or len(body_text) < 10:
        return None
    if len(heading_text) > 80 or heading_text[-1] in (".", ";", ","):
        return None
    if not any(bold for _, bold, _ in before):
        return None
    if all(bold for _, bold, _ in after):
        return None

    return heading_text, body_text


def _accept_all_text(element: Any) -> str:
    """Concatenated <w:t> text under `element`, excluding any inside a <w:del>
    subtree. Captures inline content-control (<w:sdt>) runs automatically, since
    they are descendant <w:t> of the paragraph."""
    parts: list[str] = []
    for node in element.iter(W_T):
        if not node.text:
            continue
        anc = node.getparent()
        inside_del = False
        while anc is not None and anc is not element:
            if anc.tag == W_DEL:
                inside_del = True
                break
            anc = anc.getparent()
        if not inside_del:
            parts.append(node.text)
    return "".join(parts)


def _paragraph_numbering(
    p: Any, numbering: _Numbering
) -> tuple[bool, int | None, int | None, int | None, int | None]:
    """(has_autonumber, list_level, num_id, abstract_num_id, outline_level).

    Numbering is read directly off w:p/w:pPr/w:numPr where present, else inherited
    from the paragraph style (w:pStyle -> styles.xml, basedOn chains followed) — the
    section-heading numbering that direct-only reads missed (DD-36). A paragraph
    counts as auto-numbered when it carries a numPr directly or resolves to a numId
    via its style; outlineLvl is metadata that travels with style-numbered headings."""
    pPr = p.find(W_PPR)
    if pPr is None:
        return False, None, None, None, None

    direct_numpr = pPr.find(W_NUMPR)
    direct_ilvl, direct_numid = (
        _numpr_levels(direct_numpr)
        if direct_numpr is not None
        else (
            None,
            None,
        )
    )

    pStyle = pPr.find(W_PSTYLE)
    style_ilvl, style_numid, style_outline = numbering.style(
        pStyle.get(W_VAL) if pStyle is not None else None
    )

    eff_ilvl = direct_ilvl if direct_ilvl is not None else style_ilvl
    eff_numid = direct_numid if direct_numid is not None else style_numid
    direct_outline = _int_val(pPr, W_OUTLINELVL)
    eff_outline = direct_outline if direct_outline is not None else style_outline
    if eff_outline == _BODY_OUTLINE_LEVEL:
        eff_outline = None

    has_auto = direct_numpr is not None or eff_numid is not None
    abstract = numbering.abstract_of(eff_numid)
    return has_auto, eff_ilvl, eff_numid, abstract, eff_outline


def _table_rows(tbl: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    for tr in tbl.findall(W_TR):
        rows.append([_norm(_accept_all_text(tc)) for tc in tr.findall(W_TC)])
    return rows


def _walk(
    container: Any,
    blocks: list[ExtractedBlock],
    in_cc: bool,
    numbering: _Numbering,
    doc_part: Any = None,
) -> None:
    """Append blocks for the direct w:p / w:tbl / w:sdt children of `container`,
    in document order, descending into block-level content controls (DD-45)."""
    for child in container:
        tag = child.tag
        if tag == W_P:
            # --- Inline/anchored image detection (w:drawing) — before the text guard ---
            # w:drawing is nested inside w:r inside w:p; search all descendants.
            # A single paragraph may contain multiple drawings (emit one block each).
            drawings_in_p = list(child.iter(W_DRAWING))
            if drawings_in_p:
                any_emitted = False
                for drawing in drawings_in_p:
                    try:
                        inline_or_anchor = drawing.find(WP_INLINE)
                        if inline_or_anchor is None:
                            inline_or_anchor = drawing.find(WP_ANCHOR)
                        cx: int | None = None
                        cy: int | None = None
                        if inline_or_anchor is not None:
                            extent = inline_or_anchor.find(WP_EXTENT)
                            if extent is not None:
                                cx_raw = extent.get("cx")
                                cy_raw = extent.get("cy")
                                cx = int(cx_raw) if cx_raw is not None else None
                                cy = int(cy_raw) if cy_raw is not None else None
                        blip_el = next(drawing.iter(A_BLIP), None)
                        if blip_el is not None and doc_part is not None:
                            r_id = blip_el.get(R_EMBED)
                            if r_id is not None and r_id in doc_part.rels:
                                img_part = doc_part.rels[r_id].target_part
                                img_bytes: bytes = img_part._blob
                                target_ref: str = doc_part.rels[r_id].target_ref
                                ext = Path(target_ref).suffix.lower()
                                mime = _MIME_BY_EXT.get(ext, "image/png")
                                blocks.append(
                                    ExtractedBlock(
                                        order=len(blocks),
                                        kind="attachment",
                                        text="",
                                        image_data=img_bytes,
                                        image_mime=mime,
                                        image_cx_emu=cx,
                                        image_cy_emu=cy,
                                        in_content_control=in_cc,
                                    )
                                )
                                any_emitted = True
                    except Exception:
                        pass
                if any_emitted:
                    continue
                # All image extractions failed — fall through to text processing.

            text = _norm(_accept_all_text(child))
            if not text:
                continue
            has_auto, level, num_id, abstract, outline = _paragraph_numbering(child, numbering)
            m = _LITERAL_NUM.match(text)
            inline = _split_at_line_break(child)
            if inline is not None:
                heading_text, body_text = inline
                mh = _LITERAL_NUM.match(heading_text)
                blocks.append(
                    ExtractedBlock(
                        order=len(blocks),
                        kind="paragraph",
                        text=heading_text,
                        has_autonumber=has_auto,
                        list_level=level,
                        num_id=num_id,
                        abstract_num_id=abstract,
                        outline_level=outline,
                        literal_prefix=mh.group(1) if mh else None,
                        in_content_control=in_cc,
                        is_bullet_list=numbering.is_bullet(abstract),
                    )
                )
                blocks.append(
                    ExtractedBlock(
                        order=len(blocks),
                        kind="paragraph",
                        text=body_text,
                        has_autonumber=False,
                        list_level=None,
                        num_id=None,
                        abstract_num_id=None,
                        outline_level=None,
                        literal_prefix=None,
                        in_content_control=in_cc,
                    )
                )
            else:
                blocks.append(
                    ExtractedBlock(
                        order=len(blocks),
                        kind="paragraph",
                        text=text,
                        has_autonumber=has_auto,
                        list_level=level,
                        num_id=num_id,
                        abstract_num_id=abstract,
                        outline_level=outline,
                        literal_prefix=m.group(1) if m else None,
                        in_content_control=in_cc,
                        is_bullet_list=numbering.is_bullet(abstract),
                    )
                )
        elif tag == W_TBL:
            rows = _table_rows(child)
            if any(any(c for c in row) for row in rows):
                blocks.append(
                    ExtractedBlock(
                        order=len(blocks), kind="table", rows=rows, in_content_control=in_cc
                    )
                )
        elif tag == W_SDT:
            content = child.find(W_SDTCONTENT)
            if content is not None:
                _walk(content, blocks, in_cc=True, numbering=numbering, doc_part=doc_part)


def _ceiling_chars(path: Path) -> int:
    """Non-space char count across every <w:t> in document.xml — the accept-all
    text ceiling our extraction must approach (deletions use <w:delText>, excluded)."""
    from lxml import etree

    with zipfile.ZipFile(path) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    return sum(len(re.sub(r"\s", "", el.text)) for el in root.iter(W_T) if el.text)


def _extracted_chars(blocks: list[ExtractedBlock]) -> int:
    total = 0
    for b in blocks:
        if b.kind == "paragraph":
            total += len(re.sub(r"\s", "", b.text))
        elif b.rows:
            total += sum(len(re.sub(r"\s", "", c)) for row in b.rows for c in row)
    return total


def count_tracked_changes(path: str | Path) -> tuple[int, int]:
    """(<w:ins> count, <w:del> count) across document.xml — the clean-document
    guard's tracked-change signal (DD-46). Extraction already flattens these to
    their accepted state (insertions kept, deletions dropped); this just surfaces
    that any were present, so the operator can re-upload a clean draft if needed."""
    from lxml import etree

    with zipfile.ZipFile(Path(path)) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    insertions = sum(1 for _ in root.iter(W_INS))
    deletions = sum(1 for _ in root.iter(W_DEL))
    return insertions, deletions


def _load_numbering(path: Path) -> _Numbering:
    """Build the style/abstractNum resolution context from styles.xml +
    numbering.xml. Both parts may be absent (synthetic docs); resolution then
    degrades to direct-numPr-only with empty maps."""
    from lxml import etree

    style_index: dict[str, _StyleNum] = {}
    abstract_map: dict[int, int] = {}
    bullet_abs: set[int] = set()
    with zipfile.ZipFile(path) as z:
        names = set(z.namelist())
        if "word/styles.xml" in names:
            style_index = _build_style_index(etree.fromstring(z.read("word/styles.xml")))
        if "word/numbering.xml" in names:
            numbering_root = etree.fromstring(z.read("word/numbering.xml"))
            abstract_map = _build_abstract_map(numbering_root)
            bullet_abs = _build_bullet_set(numbering_root)
    return _Numbering(style_index, abstract_map, bullet_abs)


def read_docx(path: str | Path) -> ParsedDocument:
    path = Path(path)
    doc = Document(str(path))
    body = doc.element.body
    numbering = _load_numbering(path)
    blocks: list[ExtractedBlock] = []
    _walk(body, blocks, in_cc=False, numbering=numbering, doc_part=doc.part)
    return ParsedDocument(
        blocks=blocks,
        extracted_chars=_extracted_chars(blocks),
        ceiling_chars=_ceiling_chars(path),
    )
