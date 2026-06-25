"""Issue-list export (F31, DD-60) — unresolved issues → a counterparty-safe .docx table.

Renders the contract's open negotiation points as a single Word table for principal
briefing, a counterparty screen-share, or the operator's own record (§9, §12). One
artifact, three audiences — so it is counterparty-safe by construction: only
`status='open'` issues are included (`closed` excluded, DD-65), and it carries no DB
ids, no comment threads, no internal/authority flags, and no Donna attribution
(DD-60). `donna`-initiated issues that survived to this list were operator-kept, so
they read as our point to raise → "Us".

`build_export` is a pure function of the loaded issues plus the contract node tree;
the route owns the async DB read. Clause numbers reuse the export renderer's single
numbering path (`_plan`, DD-43 / DD-02) so an issue's clause reference matches the
number the same clause carries in the contract .docx export — one numbering source,
never re-derived here.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document

from backend.models.imports import StoredNode
from backend.models.issue_export import IssueListExport, IssueRow
from backend.models.issues import StoredIssue
from backend.services.export.render_docx import _plan

UNRESOLVED_STATUSES = ("open",)

EM_DASH = "—"

_STATUS_LABEL = {"open": "Open"}
_RAISED_BY = {"operator": "Us", "counterparty": "Them", "donna": "Us"}

COLUMN_HEADERS = (
    "#",
    "Clause",
    "Issue",
    "Status",
    "Raised by",
    "Our position",
    "Their position",
    "Proposed resolution",
)

_DOCX_TABLE_STYLE = "Table Grid"
_FREE_FLOATING_SEPARATOR = "Contract-level issues (not tied to a clause)"


def _status_label(status: str) -> str:
    return _STATUS_LABEL.get(status, status)


def _raised_by(initiator: str) -> str:
    return _RAISED_BY.get(initiator, EM_DASH)


def _proposed_resolution(issue: StoredIssue) -> str:
    return issue.recommended_position or issue.donna_counter_language or EM_DASH


def _or_dash(value: str | None) -> str:
    return value if value else EM_DASH


def _positions_and_numbers(nodes: list[StoredNode]) -> tuple[dict[str, int], dict[str, str]]:
    """node_id → document-order index (every node) and node_id → derived clause
    number (clause nodes only). Both come off the renderer's single `_plan`, so the
    order tie-break and the displayed clause ref pivot on the same numbering."""
    plan = _plan(nodes)
    positions = {node.id: i for i, (node, _) in enumerate(plan)}
    numbers = {node.id: number for node, number in plan if number is not None}
    return positions, numbers


def _map_row(issue: StoredIssue, clause_numbers: dict[str, str], sequence: int) -> IssueRow:
    # `#` is the 1..n render-order reference (DD-61) so a reader can cite "item 3";
    # the raw `priority` drives the sort upstream but is an internal triage number
    # that is never printed — it is not counterparty-safe.
    clause = clause_numbers.get(issue.node_id, EM_DASH) if issue.node_id is not None else EM_DASH
    return IssueRow(
        number=str(sequence),
        clause=clause,
        issue=issue.title,
        status=_status_label(issue.status),
        raised_by=_raised_by(issue.initiator),
        our_position=_or_dash(issue.our_position),
        their_position=_or_dash(issue.their_position),
        proposed_resolution=_proposed_resolution(issue),
    )


def build_export(issues: list[StoredIssue], nodes: list[StoredNode]) -> IssueListExport:
    """Filter to unresolved, order by priority desc (NULLs last) with the document
    position as tie-break, and split clause-anchored from free-floating issues."""
    positions, numbers = _positions_and_numbers(nodes)
    unresolved = [i for i in issues if i.status in UNRESOLVED_STATUSES]
    anchored = [i for i in unresolved if i.node_id is not None]
    floating = [i for i in unresolved if i.node_id is None]

    # Tie-break for anchored issues is the node's document position; an issue on a
    # deleted/unknown node sorts after every live node (and renders clause "—").
    last_position = len(positions) + len(anchored) + 1

    def anchored_key(issue: StoredIssue) -> tuple[bool, int, int]:
        assert issue.node_id is not None
        return (
            issue.priority is None,
            -(issue.priority or 0),
            positions.get(issue.node_id, last_position),
        )

    def floating_key(issue: StoredIssue) -> tuple[bool, int, Any]:
        return (issue.priority is None, -(issue.priority or 0), issue.created_at)

    anchored.sort(key=anchored_key)
    floating.sort(key=floating_key)

    # A single 1..n sequence runs across the whole printed list in render order:
    # anchored rows first, then free-floating continuing the same count (not a
    # restart), so the briefing's "#" is stable end-to-end (DD-61).
    anchored_rows = [_map_row(i, numbers, seq) for seq, i in enumerate(anchored, start=1)]
    floating_rows = [
        _map_row(i, numbers, seq) for seq, i in enumerate(floating, start=len(anchored) + 1)
    ]

    return IssueListExport(anchored=anchored_rows, floating=floating_rows)


def _bold_cell(cell: Any) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True


def _add_data_row(table: Any, row: IssueRow) -> None:
    cells = table.add_row().cells
    values = (
        row.number,
        row.clause,
        row.issue,
        row.status,
        row.raised_by,
        row.our_position,
        row.their_position,
        row.proposed_resolution,
    )
    for j, value in enumerate(values):
        cells[j].text = value


def _add_separator_row(table: Any, label: str) -> None:
    cells = table.add_row().cells
    merged = cells[0]
    for cell in cells[1:]:
        merged = merged.merge(cell)
    merged.text = label
    _bold_cell(merged)


def render_issue_list_docx(contract_name: str, export: IssueListExport) -> bytes:
    """Pure CPU — the caller owns the async DB read and offloads this off the loop.
    An empty list still produces a valid header-only document (the operator may want
    the clean record)."""
    doc = Document()
    doc.add_heading(f"{contract_name} — Open Issues", level=1)

    if not export.anchored and not export.floating:
        doc.add_paragraph("No unresolved issues.")
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    table = doc.add_table(rows=1, cols=len(COLUMN_HEADERS))
    table.style = _DOCX_TABLE_STYLE
    header_cells = table.rows[0].cells
    for j, header in enumerate(COLUMN_HEADERS):
        header_cells[j].text = header
        _bold_cell(header_cells[j])

    for row in export.anchored:
        _add_data_row(table, row)

    if export.floating:
        _add_separator_row(table, _FREE_FLOATING_SEPARATOR)
        for row in export.floating:
            _add_data_row(table, row)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
