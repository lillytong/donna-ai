"""Deterministic DB → .docx renderer (F14, DD-43).

Export is a pure function `render(skeleton) → .docx`: the live node tree plus the
contract's `style_config` in, a clean Word document out — no AI, no network
(DD-43). It is the inverse of the import spine and doubles as the import verifier:
re-extracting a rendered document must recover the original content unchanged
(§2.1). This renders the CURRENT CLEAN state only — tracked-change export (F15) is
a later slice.

Reuses the import spine's primitives rather than re-deriving them (DD-43, one
numbering): clause numbers come from `import_.numbering.derive_numbers` (DD-02,
position-derived, never stored), so import and export pivot on the same numbering.

Content-integrity boundaries this renderer must hold:
  - Clause numbers are emitted as Word auto-numbering (`w:numPr` + a generated
    numbering definition) so the number lives outside the run text — re-extraction
    recovers the heading text exactly. BUT a real corpus stores some native
    enumerators *inside* the text (a JVA body that literally starts "5.2.1 …", or a
    list item "(a) …"). Auto-numbering those would double the marker ("5.2.1 5.2.1 …")
    or impose a synthetic decimal on a "(a)" item. So a node whose text already
    carries a leading enumerator renders verbatim with NO auto-number — the run
    text is unchanged either way, so the round-trip is unaffected.
  - `caps` / bold are display properties (`w:caps`, `w:b`), never a `str.upper()`
    or text mutation — the stored text stays original-case and round-trips intact
    while Word renders the house style (DD-37 / §2.1).

House style (DD-37) is applied as render rules, not read from config, because a
contract's `style_config` is commonly `{}`: section headings render bold,
appendix titles start a new centred page, defined terms and all-caps spans bold
inline. The uppercase transform is NOT a house rule — `caps` keys on the source
caps property (`style.level(ilvl).caps`) only, never on bold (issue #2): a bold
mixed-case heading stays mixed-case; a source-uppercase heading renders uppercase
(its stored text is already uppercase, or its level carries `caps: true`). A
populated `style_config` can still raise per-level font sizes on top.
"""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from backend.models.contract_tree import ParsedTree, TreeNode
from backend.models.imports import StoredNode
from backend.models.style import LevelStyle, StyleConfig
from backend.services.import_.numbering import derive_numbers

# High ids to avoid colliding with any abstractNum/num the default template ships.
_ABSTRACT_ID = 7777
_NUM_ID = 7777
_BULLET_ABSTRACT_ID = 7778
_BULLET_NUM_ID = 7778

# Numbering format per outline level for each scheme (DD-37). read_docx reads only
# ilvl + numId off w:numPr, so these affect Word's display, not the round-trip.
_MIXED_FMT = ("decimal", "decimal", "lowerLetter", "lowerRoman")

# A dotted-decimal outline prefix ("5.2 ", "5.2.1. ") — unambiguous, always a
# stored enumerator. A parenthesised alpha/roman marker ("(a) ", "(iv) ") — a list
# item's native marker. Either means the text already carries its number, so the
# renderer must not auto-number on top.
_DOTTED_PREFIX = re.compile(r"^\s*\d+(?:\.\d+)+\.?\s+")
_PAREN_PREFIX = re.compile(r"^\s*\([a-zA-Z]{1,4}\)\s*")

# Inline bold spans (house style): a leading quoted defined term ("Affiliate" means
# …) and all-caps legal connectives. The bolded all-caps spans are now a fixed
# allowlist of legal-connective words (WHEREAS, NOW, AND, …) rather than a length rule,
# so acronyms (DBO, JVA) are never bolded while short connectives (AND, NOW) always are —
# no length rule can separate same-length tokens, only an allowlist can. Case-sensitive
# so only the uppercase token matches.
_LEADING_DEFINED_TERM = re.compile(r'^\s*[“"][^”"]{1,80}[”"]')
_ALL_CAPS_SPAN = re.compile(
    r"\b(?:WHEREAS|NOW|THEREFORE|BY|AND|BETWEEN|AMONG|IN|WITNESS|WHEREOF|"
    r"WITNESSETH|RECITALS|BACKGROUND|PREAMBLE|HEREBY|HERETO|HEREOF|THEREOF)\b"
)


def _w(tag: str) -> str:
    return qn(f"w:{tag}")


def _group_children(nodes: list[StoredNode]) -> dict[str | None, list[StoredNode]]:
    """parent_id → siblings ordered by order_index. A parent_id absent from the
    node set (or null) is a root — mirrors ContractTreeResponse.from_rows."""
    present = {n.id for n in nodes}
    children: dict[str | None, list[StoredNode]] = {}
    for n in nodes:
        parent = n.parent_id if (n.parent_id is not None and n.parent_id in present) else None
        children.setdefault(parent, []).append(n)
    for siblings in children.values():
        siblings.sort(key=lambda n: n.order_index)
    return children


def _plan(nodes: list[StoredNode]) -> list[tuple[StoredNode, str | None]]:
    """Document-order (pre-order DFS) list of (node, derived_number). The number is
    None for non-clause nodes and clause nodes carry their DD-02 decimal-outline
    number, derived through the shared import numbering on an adapter tree."""
    children = _group_children(nodes)
    ordered: list[StoredNode] = []

    def dfs(parent: str | None) -> None:
        for node in children.get(parent, []):
            ordered.append(node)
            dfs(node.id)

    dfs(None)

    index_of = {node.id: i for i, node in enumerate(ordered)}
    tree_nodes: list[TreeNode] = []
    for i, node in enumerate(ordered):
        parent_id = node.parent_id if node.parent_id in index_of else None
        tree_nodes.append(
            TreeNode(
                index=i,
                parent_index=index_of[parent_id] if parent_id is not None else None,
                depth=0,
                order_index=node.order_index,
                kind="table" if node.content_type == "table" else "prose",
                text=(node.heading or node.body or ""),
                role=node.role,
            )
        )
    numbers = derive_numbers(ParsedTree(nodes=tree_nodes))
    return [(node, numbers.get(i)) for i, node in enumerate(ordered)]


def _depth_map(nodes: list[StoredNode]) -> dict[str, int]:
    """node id → nesting depth (root = 0). A fallback indent signal for a heading
    clause that carries no derived number."""
    children = _group_children(nodes)
    depth_of: dict[str, int] = {}

    def dfs(parent: str | None, depth: int) -> None:
        for node in children.get(parent, []):
            depth_of[node.id] = depth
            dfs(node.id, depth + 1)

    dfs(None, 0)
    return depth_of


def _lvl_text(ilvl: int, scheme: str) -> str:
    """Word lvlText: "%1.%2.…." — each ancestor level's counter, dot-joined."""
    return ".".join(f"%{i + 1}" for i in range(ilvl + 1)) + "."


def _num_fmt(ilvl: int, scheme: str) -> str:
    if scheme == "mixed":
        return _MIXED_FMT[ilvl] if ilvl < len(_MIXED_FMT) else "decimal"
    return "decimal"


def _inject_numbering(doc: Any, max_ilvl: int, scheme: str) -> None:
    """Add one multilevel abstractNum (levels 0..max_ilvl) and a num referencing it
    to the document's numbering part, so numbered paragraphs render real Word
    numbers — kept out of the run text (the content-integrity boundary)."""
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(_w("abstractNumId"), str(_ABSTRACT_ID))
    multi = OxmlElement("w:multiLevelType")
    multi.set(_w("val"), "multilevel")
    abstract.append(multi)
    for ilvl in range(max_ilvl + 1):
        lvl = OxmlElement("w:lvl")
        lvl.set(_w("ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(_w("val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(_w("val"), _num_fmt(ilvl, scheme))
        lvl.append(fmt)
        text = OxmlElement("w:lvlText")
        text.set(_w("val"), _lvl_text(ilvl, scheme))
        lvl.append(text)
        jc = OxmlElement("w:lvlJc")
        jc.set(_w("val"), "left")
        lvl.append(jc)
        abstract.append(lvl)
    numbering.insert(0, abstract)
    num = OxmlElement("w:num")
    num.set(_w("numId"), str(_NUM_ID))
    ref = OxmlElement("w:abstractNumId")
    ref.set(_w("val"), str(_ABSTRACT_ID))
    num.append(ref)
    numbering.append(num)


def _apply_numbering(paragraph: Any, ilvl: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(_w("val"), str(ilvl))
    numPr.append(ilvl_el)
    num_id = OxmlElement("w:numId")
    num_id.set(_w("val"), str(_NUM_ID))
    numPr.append(num_id)
    pPr.append(numPr)


def _inject_bullet_numbering(doc: Any, max_ilvl: int) -> None:
    """Inject a multilevel bullet abstractNum so list nodes export as native Word
    round bullets with correct indent at every nesting depth.  Each level carries
    explicit w:pPr/w:ind so Word doesn't collapse all levels to ilvl=0."""
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(_w("abstractNumId"), str(_BULLET_ABSTRACT_ID))
    multi = OxmlElement("w:multiLevelType")
    multi.set(_w("val"), "hybridMultilevel")
    abstract.append(multi)
    for ilvl in range(max_ilvl + 1):
        lvl = OxmlElement("w:lvl")
        lvl.set(_w("ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(_w("val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(_w("val"), "bullet")
        lvl.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(_w("val"), "•")
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(_w("val"), "left")
        lvl.append(jc)
        # w:ind in twips: ilvl=0 → 0 left indent (matches cockpit list_level=0),
        # each subsequent level steps in by 360 twips (0.25 inch); w:hanging
        # proportionally smaller at 180 twips.
        pPr_lvl = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(_w("left"), str(360 * ilvl))    # ilvl=0 → 0, ilvl=1 → 360, ilvl=2 → 720
        ind.set(_w("hanging"), "180")
        pPr_lvl.append(ind)
        lvl.append(pPr_lvl)
        abstract.append(lvl)
    numbering.insert(0, abstract)
    num = OxmlElement("w:num")
    num.set(_w("numId"), str(_BULLET_NUM_ID))
    ref = OxmlElement("w:abstractNumId")
    ref.set(_w("val"), str(_BULLET_ABSTRACT_ID))
    num.append(ref)
    numbering.append(num)


def _apply_bullet_numbering(paragraph: Any, ilvl: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(_w("val"), str(ilvl))
    numPr.append(ilvl_el)
    num_id = OxmlElement("w:numId")
    num_id.set(_w("val"), str(_BULLET_NUM_ID))
    numPr.append(num_id)
    pPr.append(numPr)


def _set_outline_level(paragraph: Any, ilvl: int) -> None:
    """Mark a heading paragraph with w:outlineLvl so Word renders the native
    collapse/expand triangle in editing view (ilvl 0 = top-level section)."""
    pPr = paragraph._p.get_or_add_pPr()
    ol = OxmlElement("w:outlineLvl")
    ol.set(_w("val"), str(min(ilvl, 8)))
    pPr.append(ol)


def _apply_indent(paragraph: Any, step_pt: int, ilvl: int) -> None:
    """Indent so a clause and its first-level sub-clause share an indent and only
    deeper levels step in (14 / 14.1 flush, 14.1.1 in one — DD-37 house style)."""
    if step_pt:
        paragraph.paragraph_format.left_indent = Pt(step_pt * max(0, ilvl - 1))


def _page_break_before(paragraph: Any) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    brk = OxmlElement("w:pageBreakBefore")
    pPr.append(brk)


def _style_run(run: Any, level: LevelStyle, body_size_pt: int, font: str) -> None:
    """Apply a configured LevelStyle to a run (used by the redline renderer, which
    styles its own ins/del runs)."""
    run.font.name = font
    run.font.bold = level.bold
    run.font.underline = level.underline
    run.font.all_caps = level.caps
    run.font.size = Pt(level.font_size_pt if level.font_size_pt is not None else body_size_pt)


def _run(paragraph: Any, text: str, font: str, size_pt: int, *, bold: bool, caps: bool) -> None:
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # DD-37: display-only uppercase. The run text stays original-case so the
    # round-trip recovers it; Word renders it uppercase.
    run.font.all_caps = caps


def _bold_spans(text: str) -> list[tuple[int, int]]:
    """Non-overlapping [start, end) spans of `text` to render bold inline: a leading
    quoted defined term, plus every all-caps emphasis run. Sorted, overlaps merged."""
    spans: list[tuple[int, int]] = []
    term = _LEADING_DEFINED_TERM.match(text)
    if term is not None:
        spans.append((text.index(term.group().lstrip()[0], term.start()), term.end()))
    for m in _ALL_CAPS_SPAN.finditer(text):
        spans.append((m.start(), m.end()))
    spans.sort()
    merged: list[tuple[int, int]] = []
    for start, end in spans:
        if merged and start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    return merged


def _emit_body(paragraph: Any, text: str, font: str, size_pt: int) -> None:
    """Add `text` as runs, bolding inline emphasis spans (defined terms, all-caps)
    while the rest stays regular."""
    spans = _bold_spans(text)
    if not spans:
        _run(paragraph, text, font, size_pt, bold=False, caps=False)
        return
    cursor = 0
    for start, end in spans:
        if start > cursor:
            _run(paragraph, text[cursor:start], font, size_pt, bold=False, caps=False)
        _run(paragraph, text[start:end], font, size_pt, bold=True, caps=False)
        cursor = end
    if cursor < len(text):
        _run(paragraph, text[cursor:], font, size_pt, bold=False, caps=False)


def _add_image(
    doc: Any,
    data: bytes,
    mime_type: str,
    cx_emu: int | None,
    cy_emu: int | None,
) -> None:
    """Embed image bytes into the document at original EMU dimensions."""
    import io as _io

    from docx.shared import Emu

    stream = _io.BytesIO(data)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    if cx_emu and cy_emu:
        run.add_picture(stream, width=Emu(cx_emu), height=Emu(cy_emu))
    elif cx_emu:
        run.add_picture(stream, width=Emu(cx_emu))
    else:
        run.add_picture(stream)


def _add_table(doc: Any, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            table.rows[i].cells[j].text = cell_text


def _ilvl_of(number: str) -> int:
    """Outline depth a clause number implies: "1" → 0, "5.2" → 1, "5.2.1" → 2."""
    return number.count(".")


def _carries_enumerator(text: str, number: str | None) -> bool:
    """True if `text` already opens with its own enumerator (so auto-numbering it
    would double the marker). Matches a dotted-decimal or "(a)" prefix, or the
    node's own derived number sitting at the head of the text."""
    if _DOTTED_PREFIX.match(text) or _PAREN_PREFIX.match(text):
        return True
    if number is not None:
        head = text.lstrip()
        if head.startswith(f"{number} ") or head.startswith(f"{number}."):
            return True
    return False


def render_contract_docx(
    nodes: list[StoredNode],
    style_config: dict[str, Any],
    node_images: dict[str, tuple[bytes, str, int | None, int | None]] | None = None,
) -> bytes:
    """Render the live node tree to a clean .docx (current state, no tracked
    changes). Pure CPU — the caller owns the async DB read and offloads this."""
    style = StyleConfig.from_config(style_config)
    plan = _plan(nodes)
    depth_of = _depth_map(nodes)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = style.font
    normal.font.size = Pt(style.body_font_size_pt)

    max_ilvl = max(
        (_ilvl_of(num) for node, num in plan if node.role == "clause" and num is not None),
        default=0,
    )
    _inject_numbering(doc, max_ilvl, style.numbering_scheme)

    max_list_ilvl = max(
        (max(0, depth_of[node.id] - 1) for node, _ in plan if node.content_type == "list"),
        default=0,
    )
    _inject_bullet_numbering(doc, max_list_ilvl)

    for node, number in plan:
        depth = depth_of[node.id]
        if node.content_type == "attachment":
            if node_images and node.id in node_images:
                data, mime, cx, cy = node_images[node.id]
                _add_image(doc, data, mime, cx, cy)
            continue
        if node.content_type == "table":
            _add_table(doc, node.table_data or [])
            continue
        text = node.heading if node.heading is not None else (node.body or "")
        if not text:
            continue

        is_heading = node.heading is not None
        paragraph = doc.add_paragraph()

        # Appendix title: its own centred page (DD-37 house style).
        if node.role == "appendix_title":
            _page_break_before(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caps = style.level(min(depth, len(_MIXED_FMT) - 1)).caps
            _run(paragraph, text, style.font, style.body_font_size_pt, bold=True, caps=caps)
            continue

        # Contract title (front matter): centred + larger + bold, mirroring the
        # appendix-title alignment but WITHOUT a page break — it is the first block
        # in the document. Display-only: the stored text is unchanged.
        if node.role == "title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(paragraph, text, style.font, style.title_font_size_pt, bold=True, caps=False)
            continue

        if node.content_type == "list":
            list_level = max(0, depth - 1)
            _apply_bullet_numbering(paragraph, list_level)
            _emit_body(paragraph, text, style.font, style.body_font_size_pt)
            continue

        is_clause = node.role == "clause" and number is not None
        auto_number = is_clause and not _carries_enumerator(text, number or "")

        # Section / appendix heading: bold (house style) + an uppercase transform
        # ONLY when the source caps property says so (DD-37, issue #2) — never
        # inferred from bold. A bold mixed-case heading ("6.1.1 Fees") stays
        # mixed-case; a genuinely-uppercase source heading renders uppercase.
        if is_heading:
            ilvl = _ilvl_of(number) if number is not None else min(depth, len(_MIXED_FMT) - 1)
            level = style.level(ilvl)
            size = level.font_size_pt or style.body_font_size_pt
            if auto_number:
                _apply_numbering(paragraph, ilvl)
            if number is not None:  # only indent clause headings, not appendix headings
                _apply_indent(paragraph, style.indent_per_level_pt, ilvl)
            _set_outline_level(paragraph, ilvl)
            _run(paragraph, text, style.font, size, bold=True, caps=level.caps)
            continue

        # Body paragraph: number only a clause whose text doesn't already carry one;
        # indent so a clause and its first-level sub-clause share an indent and only
        # deeper levels step in (14 / 14.1 flush; 14.1.1 in one — DD-37).
        if is_clause and number is not None:
            ilvl = _ilvl_of(number)
            size = style.level(ilvl).font_size_pt or style.body_font_size_pt
            if auto_number:
                _apply_numbering(paragraph, ilvl)
            _apply_indent(paragraph, style.indent_per_level_pt, ilvl)
            _set_outline_level(paragraph, ilvl)
            _emit_body(paragraph, text, style.font, size)
            continue

        # Appendix body paragraph: flat (no indent) to match the import/cockpit view.
        if node.role == "appendix":
            _emit_body(paragraph, text, style.font, style.body_font_size_pt)
            continue

        _emit_body(paragraph, text, style.font, style.body_font_size_pt)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
