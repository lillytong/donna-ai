"""Tracked-changes (redline) .docx renderer (F15, DD-51).

Reuses the clean renderer's numbering + styling primitives (DD-43, one numbering)
and weaves Word tracked-change runs over the diff:

  - inserted / edited-new text → `<w:ins>` wrapping `<w:r><w:t>`.
  - deleted / edited-old text  → `<w:del>` wrapping `<w:r><w:delText>` (delText, not t).
  - table insert/delete        → the whole table, each row's `<w:trPr>` carrying a
    `<w:ins>` / `<w:del>` row-property marker and the cell runs marked accordingly.
  - MOVE (structural, DD-13)    → the del+ins fallback, NOT Word's native `w:move*`:
    the node is struck (`w:del`) at its baseline position (moved-from) and inserted
    (`w:ins`) at its current position (moved-to). Native `w:moveFrom`/`w:moveTo` with
    a matched `w:name` was the PREFERRED markup (DD-51), but it requires four range
    markers (`w:moveFrom/ToRangeStart|End`) plus run wrappers whose exact OOXML could
    not be validated against Word in this environment; emitting unverifiable
    move-range XML risks a corrupt part, so the renderer falls back to the already
    Word-validated `w:ins`/`w:del` primitives (DD-51). A reviewer reads it as a move;
    Word does not label it "moved" — the flagged fidelity gap.

Each change run carries `w:author` (the operator org from config, F25/DD-44 —
never "Donna") and a `w:date`. Pure CPU given its inputs — the caller owns the
async DB reads and offloads this.

Content-integrity boundaries held (§2.1/§2.4): clause numbers stay Word
auto-numbering (outside the run text); changed text lives in `w:t` / `w:delText`
runs verbatim; unchanged nodes render exactly as the clean export (no markup), so
accepting every change reproduces the current clean copy. Deleted / moved-from nodes
are rendered struck in their baseline position and carry NO live number (they are
out of the current numbering — a renumber-only shift is not itself a change).
"""

from __future__ import annotations

import io
import itertools
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt

from backend.models.imports import StoredNode
from backend.models.redline import DeletedNode, MovedNode, NodeDiff
from backend.models.style import LevelStyle, StyleConfig
from backend.services.export.render_docx import (
    _add_table,
    _apply_numbering,
    _ilvl_of,
    _inject_numbering,
    _plan,
    _style_run,
)

# Full XML namespace (qn's nsmap has no 'xml' prefix) — preserve significant
# leading/trailing whitespace in tracked-change text runs.
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _change_element(kind: str, text: str, author: str, date: str, change_id: int) -> Any:
    """Build a `w:ins` or `w:del` element wrapping one run. Deletions use
    `w:delText` (DD-51); insertions use `w:t`."""
    wrapper = OxmlElement(f"w:{kind}")
    wrapper.set(f"{_W_NS}id", str(change_id))
    wrapper.set(f"{_W_NS}author", author)
    wrapper.set(f"{_W_NS}date", date)
    run = OxmlElement("w:r")
    text_tag = "delText" if kind == "del" else "t"
    text_el = OxmlElement(f"w:{text_tag}")
    text_el.set(_XML_SPACE, "preserve")
    text_el.text = text
    run.append(text_el)
    wrapper.append(run)
    return wrapper


def _weave(
    live_nodes: list[StoredNode],
    deleted_nodes: list[DeletedNode],
    moved: dict[str, MovedNode],
) -> list[tuple[str, Any]]:
    """Document-order (pre-order DFS) weave of three entry kinds, each at its own
    position: live nodes (current position), deleted nodes (baseline position), and
    moved-from ghosts (baseline position). Returns (kind, payload) pairs where kind
    is "live" | "deleted" | "moved_from". A moved node appears twice — live at its
    new position and a moved-from ghost at its old one. Numbering still derives from
    the live-only tree; this only fixes WHERE struck entries appear."""
    real_ids = {n.id for n in live_nodes} | {d.id for d in deleted_nodes}

    def norm(parent_id: str | None) -> str | None:
        return parent_id if (parent_id is not None and parent_id in real_ids) else None

    # (key, parent_key, order_index, kind, payload). Ghost keys are synthetic so a
    # moved node's two placements never collide; ghosts attach under their baseline
    # parent (a real id) and are leaves (no recursion).
    entries: list[tuple[str, str | None, int, str, Any]] = []
    for n in live_nodes:
        entries.append((n.id, norm(n.parent_id), n.order_index, "live", n))
    for d in deleted_nodes:
        entries.append((d.id, norm(d.parent_id), d.order_index, "deleted", d))
    for m in moved.values():
        entries.append(
            (f"mf::{m.id}", norm(m.baseline_parent_id), m.baseline_order_index, "moved_from", m)
        )

    children: dict[str | None, list[tuple[str, str | None, int, str, Any]]] = {}
    for e in entries:
        children.setdefault(e[1], []).append(e)
    for siblings in children.values():
        siblings.sort(key=lambda e: e[2])

    ordered: list[tuple[str, Any]] = []

    def dfs(parent_key: str | None) -> None:
        for e in children.get(parent_key, []):
            ordered.append((e[3], e[4]))
            if e[3] != "moved_from":
                dfs(e[0])

    dfs(None)
    return ordered


def _apply_clause_format(
    paragraph: Any, node: StoredNode, number: str | None, style: StyleConfig
) -> LevelStyle:
    if node.role == "clause" and number is not None:
        ilvl = _ilvl_of(number)
        _apply_numbering(paragraph, ilvl)
        if style.indent_per_level_pt:
            paragraph.paragraph_format.left_indent = Pt(style.indent_per_level_pt * ilvl)
        return style.level(ilvl)
    return LevelStyle()


def _render_unchanged(doc: Any, node: StoredNode, number: str | None, style: StyleConfig) -> None:
    text = node.heading if node.heading is not None else (node.body or "")
    if not text:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    level = _apply_clause_format(paragraph, node, number, style)
    _style_run(run, level, style.body_font_size_pt, style.font)


def _render_inserted(
    doc: Any,
    node: StoredNode,
    number: str | None,
    style: StyleConfig,
    diff: NodeDiff,
    author: str,
    date: str,
    ids: Any,
) -> None:
    text = diff.text_after or ""
    if not text:
        return
    paragraph = doc.add_paragraph()
    _apply_clause_format(paragraph, node, number, style)
    paragraph._p.append(_change_element("ins", text, author, date, next(ids)))


def _render_edited(
    doc: Any,
    node: StoredNode,
    number: str | None,
    style: StyleConfig,
    diff: NodeDiff,
    author: str,
    date: str,
    ids: Any,
) -> None:
    paragraph = doc.add_paragraph()
    _apply_clause_format(paragraph, node, number, style)
    before = diff.text_before or ""
    after = diff.text_after or ""
    if before:
        paragraph._p.append(_change_element("del", before, author, date, next(ids)))
    if after:
        paragraph._p.append(_change_element("ins", after, author, date, next(ids)))


def _row_change_marker(row: Any, kind: str, author: str, date: str, change_id: int) -> None:
    """Stamp a `<w:trPr>` `<w:ins>` / `<w:del>` row-property marker — Word's
    inserted/deleted-table-row tracked change (DD-51)."""
    trPr = row._tr.get_or_add_trPr()
    marker = OxmlElement(f"w:{kind}")
    marker.set(f"{_W_NS}id", str(change_id))
    marker.set(f"{_W_NS}author", author)
    marker.set(f"{_W_NS}date", date)
    trPr.append(marker)


def _add_table_tracked(
    doc: Any, rows: list[list[str]], kind: str, author: str, date: str, ids: Any
) -> None:
    """Render a whole table marked inserted ("ins") or deleted ("del"): every row
    carries a trPr ins/del marker and every cell's text is a tracked run. A deleted
    table whose cell content is unrecoverable (snapshot stores no table_data) is
    struck as a single empty cell — the flagged fidelity gap."""
    if not rows:
        if kind != "del":
            return
        rows = [[""]]
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        _row_change_marker(table.rows[i], kind, author, date, next(ids))
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ""
            paragraph = table.rows[i].cells[j].paragraphs[0]
            paragraph._p.append(_change_element(kind, cell_text, author, date, next(ids)))


def _render_deleted(doc: Any, node: DeletedNode, author: str, date: str, ids: Any) -> None:
    if node.content_type == "table":
        _add_table_tracked(doc, [], "del", author, date, ids)
        return
    if not node.text:
        return
    paragraph = doc.add_paragraph()
    paragraph._p.append(_change_element("del", node.text, author, date, next(ids)))


def _render_moved_from(doc: Any, moved: MovedNode, author: str, date: str, ids: Any) -> None:
    """The struck moved-from ghost at the node's baseline position."""
    if moved.content_type == "table":
        _add_table_tracked(doc, moved.table_data or [], "del", author, date, ids)
        return
    if not moved.baseline_text:
        return
    paragraph = doc.add_paragraph()
    paragraph._p.append(_change_element("del", moved.baseline_text, author, date, next(ids)))


def _render_moved_to(
    doc: Any,
    node: StoredNode,
    number: str | None,
    style: StyleConfig,
    moved: MovedNode,
    author: str,
    date: str,
    ids: Any,
) -> None:
    """The inserted moved-to node at its current position (carries its new number)."""
    if node.content_type == "table":
        _add_table_tracked(doc, node.table_data or [], "ins", author, date, ids)
        return
    text = moved.current_text
    if not text:
        return
    paragraph = doc.add_paragraph()
    _apply_clause_format(paragraph, node, number, style)
    paragraph._p.append(_change_element("ins", text, author, date, next(ids)))


def render_redline_docx(
    live_nodes: list[StoredNode],
    diffs: dict[str, NodeDiff],
    deleted_nodes: list[DeletedNode],
    style_config: dict[str, Any],
    author: str,
    timestamp: str,
    moved: dict[str, MovedNode] | None = None,
    inserted_tables: set[str] | None = None,
) -> bytes:
    """Render the working copy against its baseline as a tracked-changes .docx.

    `diffs` keys live node ids (inserted/edited); `deleted_nodes` carry their own
    baseline position + struck text. `moved` keys live node ids (rendered as a
    del+ins move); `inserted_tables` is the set of live table-node ids to mark
    inserted. Pure CPU — the caller offloads it."""
    moved = moved or {}
    inserted_tables = inserted_tables or set()
    style = StyleConfig.from_config(style_config)
    plan = _plan(live_nodes)
    numbers = {node.id: num for node, num in plan}

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = style.font
    normal.font.size = Pt(style.body_font_size_pt)

    max_ilvl = max(
        (_ilvl_of(num) for node, num in plan if node.role == "clause" and num is not None),
        default=0,
    )
    _inject_numbering(doc, max_ilvl, style.numbering_scheme)

    ids = itertools.count(1)
    for kind, payload in _weave(live_nodes, deleted_nodes, moved):
        if kind == "deleted":
            _render_deleted(doc, payload, author, timestamp, ids)
            continue
        if kind == "moved_from":
            _render_moved_from(doc, payload, author, timestamp, ids)
            continue
        node = payload
        number = numbers.get(node.id)
        if node.id in moved:
            _render_moved_to(doc, node, number, style, moved[node.id], author, timestamp, ids)
            continue
        if node.content_type == "table":
            if node.id in inserted_tables:
                _add_table_tracked(doc, node.table_data or [], "ins", author, timestamp, ids)
            else:
                _add_table(doc, node.table_data or [])
            continue
        diff = diffs.get(node.id)
        if diff is None:
            _render_unchanged(doc, node, number, style)
        elif diff.change_type == "inserted":
            _render_inserted(doc, node, number, style, diff, author, timestamp, ids)
        else:
            _render_edited(doc, node, number, style, diff, author, timestamp, ids)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
